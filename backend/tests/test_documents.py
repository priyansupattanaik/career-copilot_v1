from io import BytesIO

from docx import Document

from app.documents import extract_sections, extract_text, safe_filename, validate_document


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python, SQL")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_is_validated_and_extracted():
    content = make_docx()
    mime = validate_document("resume.docx", None, content, 10 * 1024 * 1024)
    text = extract_text(content, mime)
    sections = extract_sections(text)
    assert "Python" in text
    assert sections["sections"]["skills"] == ["Python, SQL"]


def test_filename_is_sanitized():
    assert safe_filename("../../my resume.docx") == "my_resume.docx"
