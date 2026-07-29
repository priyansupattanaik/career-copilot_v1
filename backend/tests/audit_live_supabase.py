import asyncio
import io
import json
import logging
import uuid

import httpx
from docx import Document
from supabase import ClientOptions, create_client

from app.config import get_settings
from app.main import app
from app.supabase_clients import create_admin_supabase_client

OWNED_TABLES = [
    "profiles",
    "candidate_preferences",
    "candidate_skills",
    "candidate_experiences",
    "candidate_projects",
    "candidate_education",
    "candidate_certifications",
    "candidate_languages",
    "candidate_links",
    "resumes",
    "resume_versions",
    "job_descriptions",
    "ats_analyses",
    "ats_evidence",
    "resume_suggestions",
    "resume_exports",
    "interview_sessions",
    "interview_questions",
    "interview_responses",
    "interview_reports",
    "learning_paths",
    "learning_items",
    "learning_resources",
    "job_recommendations",
    "saved_jobs",
    "notification_preferences",
    "privacy_preferences",
    "activity_events",
    "user_notifications",
]
BUCKET_FIXTURES = {
    "candidate-documents": ("audit.pdf", b"%PDF-1.4\n%%EOF", "application/pdf"),
    "candidate-avatars": ("audit.png", b"\x89PNG\r\n\x1a\nAUDIT", "image/png"),
    "interview-media": ("audit.webm", b"\x1aE\xdf\xa3AUDIT", "video/webm"),
}


def user_client(settings, token):
    client = create_client(
        settings.supabase_url,
        settings.supabase_publishable_key,
        options=ClientOptions(headers={"Authorization": f"Bearer {token}"}),
    )
    client.postgrest.auth(token)
    return client


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python, SQL, accessibility")
    document.add_heading("Experience", level=1)
    document.add_paragraph("Built evidence-led product workflows.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def make_pdf() -> bytes:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        (
            b"<< /Length 70 >>\nstream\nBT /F1 12 Tf 72 720 Td "
            b"(Skills Python SQL accessibility) Tj ET\nendstream"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode())
        content.extend(body)
        content.extend(b"\nendobj\n")
    xref = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode())
    content.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return bytes(content)


async def main() -> None:
    logging.getLogger("httpx").setLevel(logging.WARNING)
    logging.getLogger("career_copilot.api").setLevel(logging.WARNING)
    settings = get_settings()
    admin = create_admin_supabase_client(settings)
    anonymous = create_client(settings.supabase_url, settings.supabase_publishable_key)
    suffix = uuid.uuid4().hex
    credentials = [
        (f"audit_user_a_{suffix}@example.invalid", f"Audit-{suffix}-A1!"),
        (f"audit_user_b_{suffix}@example.invalid", f"Audit-{suffix}-B1!"),
        (f"audit_delete_{suffix}@example.invalid", f"Audit-{suffix}-C1!"),
    ]
    user_ids: list[str] = []
    storage_cleanup: dict[str, list[str]] = {bucket: [] for bucket in BUCKET_FIXTURES}
    job_id = ""
    summary: dict[str, object] = {}
    try:
        clients = []
        tokens = []
        for email, password in credentials:
            created = admin.auth.admin.create_user(
                {"email": email, "password": password, "email_confirm": True}
            )
            user_ids.append(str(created.user.id))
            browser = create_client(settings.supabase_url, settings.supabase_publishable_key)
            signed_in = browser.auth.sign_in_with_password({"email": email, "password": password})
            tokens.append(signed_in.session.access_token)
            clients.append(user_client(settings, signed_in.session.access_token))

        a, b, _ = clients
        a_id, b_id, delete_id = user_ids
        for table in [
            "profiles",
            "candidate_preferences",
            "notification_preferences",
            "privacy_preferences",
        ]:
            owner_column = "id" if table == "profiles" else "user_id"
            assert a.table(table).select(owner_column).eq(owner_column, a_id).single().execute().data
        summary["profile_trigger"] = "4/4 defaults"

        transport = httpx.ASGITransport(app=app, raise_app_exceptions=False)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as api:
            assert (await api.get("/api/v1/profile")).status_code == 401
            assert (
                await api.get("/api/v1/profile", headers={"Authorization": "Token invalid"})
            ).status_code == 401
            assert (
                await api.get("/api/v1/profile", headers={"Authorization": "Bearer invalid"})
            ).status_code == 401
            headers = {"Authorization": f"Bearer {tokens[0]}"}
            b_headers = {"Authorization": f"Bearer {tokens[1]}"}

            profile = await api.patch(
                "/api/v1/profile",
                headers=headers,
                json={
                    "full_name": "Audit User A",
                    "headline": "Evidence Engineer",
                    "phone": "+91 9000000000",
                    "location": "Pune",
                    "current_role": "Analyst",
                    "years_experience": 4,
                    "career_level": "mid",
                    "career_goal": "Build secure career systems",
                },
            )
            assert profile.status_code == 200, profile.text
            preferences = await api.put(
                "/api/v1/profile/preferences",
                headers=headers,
                json={
                    "target_roles": ["Evidence Engineer"],
                    "preferred_industries": ["Technology"],
                    "preferred_locations": ["Pune"],
                    "work_modes": ["hybrid"],
                    "employment_types": ["full_time"],
                    "notice_period_days": 30,
                    "willing_to_relocate": False,
                    "work_authorization": "India",
                    "salary_min": 1000000,
                    "salary_max": 1500000,
                    "salary_currency": "INR",
                },
            )
            assert preferences.status_code == 200, preferences.text

            resources = {
                "skills": {"name": "Python", "normalized_name": f"python-{suffix}"},
                "experiences": {"company_name": "Audit Co", "role_title": "Engineer"},
                "projects": {"title": "Audit Project", "skills": ["Python"]},
                "education": {"institution": "Audit University", "degree": "BTech"},
                "certifications": {"name": "Audit Certificate", "issuer": "Audit"},
                "languages": {"language": "English", "normalized_language": f"english-{suffix}"},
                "links": {"link_type": "website", "url": "https://example.invalid/audit"},
            }
            for resource, payload in resources.items():
                response = await api.post(f"/api/v1/profile/{resource}", headers=headers, json=payload)
                assert response.status_code == 201, response.text

            docx = make_docx()
            pdf = make_pdf()
            resume_docx = await api.post(
                "/api/v1/resumes",
                headers=headers,
                data={"title": "Audit DOCX resume"},
                files={
                    "file": (
                        "audit.docx",
                        docx,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            assert resume_docx.status_code == 201, resume_docx.text
            resume_docx_data = resume_docx.json()
            storage_cleanup[settings.document_bucket].append(resume_docx_data["version"]["storage_path"])
            resume_id = resume_docx_data["resume"]["id"]
            version_id = resume_docx_data["version"]["id"]
            assert resume_docx_data["version"]["plain_text"]
            assert len(resume_docx_data["version"]["sha256"]) == 64

            resume_pdf = await api.post(
                "/api/v1/resumes",
                headers=headers,
                data={"title": "Audit PDF resume"},
                files={"file": ("audit.pdf", pdf, "application/pdf")},
            )
            assert resume_pdf.status_code == 201, resume_pdf.text
            resume_pdf_data = resume_pdf.json()
            storage_cleanup[settings.document_bucket].append(resume_pdf_data["version"]["storage_path"])

            correction = {"schema_version": "resume-extraction-v1", "corrections": {"headline": "Corrected"}}
            patched = await api.patch(
                f"/api/v1/resume-versions/{version_id}/extraction",
                headers=headers,
                json={"structured_content": correction},
            )
            assert patched.status_code == 200 and patched.json()["structured_content"] == correction
            assert (
                await api.post(f"/api/v1/resume-versions/{version_id}/confirm", headers=headers)
            ).status_code == 200
            assert (await api.get(f"/api/v1/resume-versions/{version_id}", headers=headers)).json()[
                "structured_content"
            ] == correction
            assert (await api.get(f"/api/v1/resumes/{resume_id}", headers=b_headers)).status_code == 404

            invalid_files = [
                ("empty.pdf", b"", "application/pdf", 400),
                ("wrong.txt", b"text", "text/plain", 415),
                ("spoofed.pdf", b"not pdf", "application/pdf", 415),
                (
                    "corrupt.docx",
                    b"not a zip",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    415,
                ),
                ("oversized.pdf", b"%PDF-" + b"x" * settings.document_max_bytes, "application/pdf", 413),
            ]
            for filename, content, mime, expected in invalid_files:
                result = await api.post(
                    f"/api/v1/resumes/{resume_id}/versions",
                    headers=headers,
                    files={"file": (filename, content, mime)},
                )
                assert result.status_code == expected, (filename, result.status_code, result.text)

            jd_text = await api.post(
                "/api/v1/job-descriptions",
                headers=headers,
                json={
                    "title": "Audit text JD",
                    "company": "Audit Co",
                    "role_title": "Evidence Engineer",
                    "raw_text": "Requirements\nPython\nSQL\nSecure persistence and accessibility experience.",
                },
            )
            assert jd_text.status_code == 201, jd_text.text
            jd_text_id = jd_text.json()["id"]
            jd_correction = {
                "schema_version": "jd-extraction-v1",
                "corrections": {"role": "Evidence Engineer"},
            }
            assert (
                await api.patch(
                    f"/api/v1/job-descriptions/{jd_text_id}/extraction",
                    headers=headers,
                    json={"structured_content": jd_correction},
                )
            ).status_code == 200
            assert (
                await api.post(f"/api/v1/job-descriptions/{jd_text_id}/confirm", headers=headers)
            ).status_code == 200

            for filename, content, mime in [
                ("audit-jd.pdf", pdf, "application/pdf"),
                (
                    "audit-jd.docx",
                    docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ]:
                uploaded = await api.post(
                    "/api/v1/job-descriptions/upload",
                    headers=headers,
                    data={"title": filename, "company": "Audit Co", "role_title": "Engineer"},
                    files={"file": (filename, content, mime)},
                )
                assert uploaded.status_code == 201, uploaded.text
                storage_cleanup[settings.document_bucket].append(uploaded.json()["storage_path"])

            analysis = (
                a.table("ats_analyses")
                .insert(
                    {
                        "user_id": a_id,
                        "resume_version_id": version_id,
                        "job_description_id": jd_text_id,
                        "status": "draft",
                    }
                )
                .execute()
                .data[0]
            )
            evidence = (
                a.table("ats_evidence")
                .insert(
                    {
                        "user_id": a_id,
                        "analysis_id": analysis["id"],
                        "requirement_text": "Python",
                        "match_status": "unverified",
                    }
                )
                .execute()
                .data[0]
            )
            a.table("resume_suggestions").insert(
                {
                    "user_id": a_id,
                    "analysis_id": analysis["id"],
                    "resume_version_id": version_id,
                    "section_key": "skills",
                    "original_text": "Python",
                    "suggested_text": "Python",
                    "supporting_evidence_ids": [evidence["id"]],
                }
            ).execute()
            a.table("resume_exports").insert(
                {
                    "user_id": a_id,
                    "resume_version_id": version_id,
                    "export_format": "pdf",
                    "storage_path": f"{a_id}/exports/audit.pdf",
                }
            ).execute()
            assert (
                await api.get(f"/api/v1/ats-analyses/{analysis['id']}", headers=headers)
            ).status_code == 200
            assert (
                await api.get(f"/api/v1/ats-analyses/{analysis['id']}/suggestions", headers=headers)
            ).status_code == 200
            assert (await api.post("/api/v1/ats-analyses", headers=headers, json={})).status_code == 503

            interview = await api.post(
                "/api/v1/interviews",
                headers=headers,
                json={"mode": "behavioural", "question_count": 1, "duration_minutes": 5},
            )
            assert interview.status_code == 201, interview.text
            session_id = interview.json()["id"]
            question = (
                a.table("interview_questions")
                .insert(
                    {
                        "user_id": a_id,
                        "session_id": session_id,
                        "position": 1,
                        "question": "Describe an evidence-led decision.",
                    }
                )
                .execute()
                .data[0]
            )
            assert (
                await api.post(f"/api/v1/interviews/{session_id}/start", headers=headers)
            ).status_code == 200
            response = await api.post(
                f"/api/v1/interviews/{session_id}/responses",
                headers=headers,
                json={
                    "question_id": question["id"],
                    "typed_response": "I verified the source records.",
                    "duration_seconds": 12,
                },
            )
            assert response.status_code == 201, response.text
            a.table("interview_reports").insert({"user_id": a_id, "session_id": session_id}).execute()
            completed = await api.post(f"/api/v1/interviews/{session_id}/complete", headers=headers)
            assert completed.status_code == 200 and completed.json()["report"] is None

            learning = await api.post(
                "/api/v1/learning-paths",
                headers=headers,
                json={
                    "title": "Audit learning path",
                    "description": "Temporary",
                    "source_type": "candidate_selected",
                },
            )
            assert learning.status_code == 201, learning.text
            path_id = learning.json()["id"]
            item = (
                a.table("learning_items")
                .insert({"user_id": a_id, "learning_path_id": path_id, "position": 1, "title": "Audit item"})
                .execute()
                .data[0]
            )
            a.table("learning_resources").insert(
                {
                    "user_id": a_id,
                    "learning_item_id": item["id"],
                    "title": "Audit resource",
                    "url": "https://example.invalid/resource",
                }
            ).execute()
            assert (await api.get(f"/api/v1/learning-paths/{path_id}", headers=headers)).status_code == 200

            job = (
                admin.table("jobs")
                .insert(
                    {
                        "external_source": "audit",
                        "external_id": suffix,
                        "title": "Audit Job",
                        "company": "Audit Co",
                        "description": "Temporary audit job",
                    }
                )
                .execute()
                .data[0]
            )
            job_id = job["id"]
            a.table("job_recommendations").insert(
                {"user_id": a_id, "job_id": job_id, "resume_version_id": version_id}
            ).execute()
            assert (await api.get(f"/api/v1/jobs/{job_id}", headers=headers)).status_code == 200
            assert (await api.post(f"/api/v1/saved-jobs/{job_id}", headers=headers)).status_code == 201
            assert (
                await api.patch(
                    f"/api/v1/saved-jobs/{job_id}",
                    headers=headers,
                    json={"status": "applied", "notes": "Audit"},
                )
            ).status_code == 200
            assert len((await api.get("/api/v1/saved-jobs", headers=headers)).json()) == 1

            notifications = await api.put(
                "/api/v1/settings/notifications",
                headers=headers,
                json={
                    "job_alerts": True,
                    "learning_reminders": True,
                    "interview_reminders": False,
                    "product_updates": False,
                    "email_frequency": "weekly",
                },
            )
            privacy = await api.put(
                "/api/v1/settings/privacy",
                headers=headers,
                json={
                    "camera_permission": "disabled",
                    "microphone_permission": "ask",
                    "recording_retention_days": 0,
                    "resume_processing_consent": True,
                    "job_recommendation_consent": False,
                    "profile_visibility": "private",
                },
            )
            assert notifications.status_code == 200 and privacy.status_code == 200
            a.table("user_notifications").insert(
                {"user_id": a_id, "notification_type": "audit", "title": "Audit", "message": "Temporary"}
            ).execute()

            summary["api_auth"] = "missing/malformed/invalid/valid verified"
            summary["api_categories"] = "profile,resume,jd,ats,interview,learning,jobs,settings"
            summary["documents"] = "PDF,DOCX,text JD and invalid-file cases"

            delete_response = await api.delete(
                "/api/v1/account",
                headers={"Authorization": f"Bearer {tokens[2]}", "X-Confirm-Delete": "DELETE MY ACCOUNT"},
            )
            assert delete_response.status_code == 204, delete_response.text
            user_ids.remove(delete_id)

        for table in OWNED_TABLES:
            owner = "id" if table == "profiles" else "user_id"
            assert a.table(table).select(owner).eq(owner, a_id).limit(1).execute().data, table
            assert b.table(table).select(owner).eq(owner, a_id).execute().data == [], table
            anonymous_denied = False
            try:
                anonymous_denied = anonymous.table(table).select(owner).eq(owner, a_id).execute().data == []
            except Exception:
                anonymous_denied = True
            assert anonymous_denied, table
            cross_update = {"full_name": "blocked"} if table == "profiles" else {"user_id": b_id}
            assert b.table(table).update(cross_update).eq(owner, a_id).execute().data == [], table
            assert b.table(table).delete().eq(owner, a_id).execute().data == [], table
        summary["rls_tables"] = f"{len(OWNED_TABLES)}/{len(OWNED_TABLES)} cross-user and anonymous denial"

        spare = (
            a.table("candidate_skills")
            .insert({"user_id": a_id, "name": "Disposable", "normalized_name": f"disposable-{suffix}"})
            .execute()
            .data[0]
        )
        assert (
            a.table("candidate_skills")
            .update({"proficiency": "verified"})
            .eq("id", spare["id"])
            .execute()
            .data
        )
        assert a.table("candidate_skills").delete().eq("id", spare["id"]).execute().data
        forged = False
        try:
            a.table("candidate_skills").insert(
                {"user_id": b_id, "name": "Forged", "normalized_name": f"forged-{suffix}"}
            ).execute()
        except Exception:
            forged = True
        assert forged
        summary["rls_crud"] = "own insert/read/update/delete and forged ownership denial"

        for bucket, (filename, content, mime) in BUCKET_FIXTURES.items():
            path = f"{a_id}/audit/{uuid.uuid4()}-{filename}"
            storage_cleanup[bucket].append(path)
            a.storage.from_(bucket).upload(path, content, {"content-type": mime, "upsert": "false"})
            assert a.storage.from_(bucket).download(path)
            a.storage.from_(bucket).update(path, content + b"2", {"content-type": mime, "upsert": "true"})
            assert b.storage.from_(bucket).list(f"{a_id}/audit") == []
            for client in (b, anonymous):
                blocked = False
                try:
                    client.storage.from_(bucket).download(path)
                except Exception:
                    blocked = True
                assert blocked, bucket
            blocked_upload = False
            try:
                b.storage.from_(bucket).upload(
                    f"{a_id}/audit/{uuid.uuid4()}-{filename}", content, {"content-type": mime}
                )
            except Exception:
                blocked_upload = True
            assert blocked_upload, bucket
            a.storage.from_(bucket).remove([path])
            storage_cleanup[bucket].remove(path)
        summary["storage"] = "3/3 private buckets own CRUD and cross-user/anonymous denial"

        print(json.dumps(summary, sort_keys=True))
    finally:
        for bucket, paths in storage_cleanup.items():
            if paths:
                try:
                    admin.storage.from_(bucket).remove(paths)
                except Exception:
                    pass
        if job_id:
            try:
                admin.table("jobs").delete().eq("id", job_id).execute()
            except Exception:
                pass
        for user_id in user_ids:
            try:
                admin.auth.admin.delete_user(user_id)
            except Exception:
                pass
        print("audit_temporary_cleanup=complete")


if __name__ == "__main__":
    asyncio.run(main())
