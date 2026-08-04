import io
import pytest
from unittest.mock import patch
from docx import Document

from app.core.config import Settings
from app.core.errors import ApiError
from app.features.document_parsing.schemas import (
    ParsedResumeSchema,
    FieldWrapper,
    ConfidenceLevel,
    ContactSection,
    SkillItem,
    ExperienceItem,
    ProjectItem,
    EducationItem,
    CertificationItem,
    LicenceItem,
    AchievementItem,
    PublicationItem,
    LanguageItem,
    VolunteerItem,
    TrainingItem,
    LinkItem,
    AdditionalSectionItem,
    UnclassifiedBlock,
    Stage1DocumentStructure,
    Stage2ContactAndSummary,
    Stage3Skills,
    Stage4Experience,
    Stage5Projects,
    Stage6EducationAndCredentials,
    create_empty_field_wrapper,
)
from app.features.document_parsing.source_blocks import SourceBlock, SourceBlockCollection
from app.features.document_parsing.extractors.pdf import parse_pdf_to_blocks, SCANNED_DENSITY_THRESHOLD_PER_PAGE
from app.features.document_parsing.extractors.docx import parse_docx_to_blocks
from app.features.document_parsing.extractors.ocr import is_ocr_available, process_scanned_pdf, ExtractionResult
from app.features.document_parsing.extractors import extract_document_blocks


# --- 1. Config & Settings Tests ---

def test_groq_resume_parser_config_defaults():
    settings = Settings(
        app_name="Test App",
        app_env="test",
        api_v1_prefix="/api/v1",
        public_api_base_url="http://localhost:8000",
        log_level="INFO",
        frontend_origins=["http://localhost:3000"],
        database_path=":memory:",
        auth_secret="secret",
        local_storage_dir="/tmp/storage",
        document_bucket="docs",
        avatar_bucket="avatars",
        interview_bucket="interviews",
        nvidia_base_url="https://api.nvidia.com",
        nvidia_model="deepseek",
        nvidia_prompt_version="v1",
        groq_api_key="",
        groq_base_url="https://api.groq.com/openai/v1",
        groq_model="llama-3.3-70b-versatile",
        llm_provider="groq",
    )
    assert settings.groq_resume_parser_enabled is True
    assert settings.groq_resume_parser_model == "openai/gpt-oss-120b"
    assert settings.groq_resume_parser_fallback_model == "llama-3.3-70b-versatile"
    assert settings.groq_resume_parser_timeout_seconds == 60.0
    assert settings.groq_resume_parser_max_retries == 2
    assert settings.groq_resume_parser_max_input_tokens == 110000
    assert settings.groq_resume_parser_temperature == 0.0
    assert settings.groq_resume_parser_configured is False  # groq_api_key is empty

    # Set GROQ_API_KEY
    settings_with_key = settings.model_copy(update={"groq_api_key": "gsk_test123"})
    assert settings_with_key.groq_resume_parser_configured is True



# --- 2. Pydantic Schemas Tests ---

def test_parsed_resume_schema_18_sections():
    schema = ParsedResumeSchema()
    assert isinstance(schema.contact, ContactSection)
    assert isinstance(schema.professional_summary, FieldWrapper)
    assert isinstance(schema.target_role, FieldWrapper)
    assert isinstance(schema.skills, list)
    assert isinstance(schema.experience, list)
    assert isinstance(schema.projects, list)
    assert isinstance(schema.education, list)
    assert isinstance(schema.certifications, list)
    assert isinstance(schema.licences, list)
    assert isinstance(schema.achievements, list)
    assert isinstance(schema.publications, list)
    assert isinstance(schema.languages, list)
    assert isinstance(schema.volunteer_experience, list)
    assert isinstance(schema.training, list)
    assert isinstance(schema.links, list)
    assert isinstance(schema.additional_sections, list)
    assert isinstance(schema.warnings, list)
    assert isinstance(schema.unclassified_blocks, list)

    json_schema = ParsedResumeSchema.model_json_schema()
    assert "properties" in json_schema
    assert "contact" in json_schema["properties"]


def test_field_wrapper_default_and_helper():
    fw = FieldWrapper[str](value="Python", evidence_block_ids=["page-1-block-01"], confidence=ConfidenceLevel.HIGH)
    assert fw.value == "Python"
    assert fw.evidence_block_ids == ["page-1-block-01"]
    assert fw.confidence == ConfidenceLevel.HIGH
    assert fw.warning is None

    empty_fw = create_empty_field_wrapper("Ambiguous value")
    assert empty_fw.value is None
    assert empty_fw.evidence_block_ids == []
    assert empty_fw.confidence == ConfidenceLevel.HIGH
    assert empty_fw.warning == "Ambiguous value"


def test_multi_stage_sub_schemas():
    stg1 = Stage1DocumentStructure(detected_sections=["skills", "experience"])
    assert stg1.detected_sections == ["skills", "experience"]

    stg2 = Stage2ContactAndSummary()
    assert isinstance(stg2.contact, ContactSection)

    stg3 = Stage3Skills(skills=[SkillItem(name=FieldWrapper(value="Python"))])
    assert len(stg3.skills) == 1
    assert stg3.skills[0].name.value == "Python"


# --- 3. SourceBlock Tests ---

def test_source_block_deterministic_creation():
    block = SourceBlock.create(
        page=1,
        order=7,
        text="Built REST APIs using Python",
        block_type="paragraph",
        heading_context="Experience",
        bounding_box=[10.0, 20.0, 300.0, 50.0],
    )
    assert block.block_id == "page-1-block-07"
    assert block.page == 1
    assert block.order == 7
    assert block.text == "Built REST APIs using Python"
    assert block.block_type == "paragraph"
    assert block.heading_context == "Experience"
    assert block.bounding_box == (10.0, 20.0, 300.0, 50.0)


def test_source_block_collection_methods():
    b1 = SourceBlock.create(page=1, order=1, text="Alice Smith", block_type="heading")
    b2 = SourceBlock.create(page=1, order=2, text="Software Engineer", heading_context="Alice Smith")
    b3 = SourceBlock.create(page=2, order=1, text="Education details")

    col = SourceBlockCollection(blocks=[b1, b2, b3])
    assert col.get_by_id("page-1-block-01") == b1
    assert col.get_by_id("page-9-block-99") is None
    assert len(col.get_by_page(1)) == 2
    assert col.total_character_count() == len("Alice Smith") + len("Software Engineer") + len("Education details")
    assert col.text_density_per_page() == {1: len("Alice Smith") + len("Software Engineer"), 2: len("Education details")}


# --- 4. PDF Extractor Tests ---

def test_pdf_extractor_pypdf_fallback():
    # Test pypdf fallback with mock content
    with patch("app.features.document_parsing.extractors.pdf._parse_pdf_fitz", side_effect=ImportError):
        with patch("app.features.document_parsing.extractors.pdf._parse_pdf_pdfplumber", side_effect=ImportError):
            # Create a simple PDF using reportlab or mock PdfReader
            from pypdf import PdfWriter
            writer = PdfWriter()
            writer.add_blank_page(width=612, height=792)
            stream = io.BytesIO()
            writer.write(stream)
            pdf_bytes = stream.getvalue()

            blocks, is_scanned = parse_pdf_to_blocks(pdf_bytes)
            # Blank PDF should trigger is_scanned=True because density < 20 chars/page
            assert is_scanned is True
            assert isinstance(blocks, list)


def test_pdf_extractor_encrypted_pdf():
    mock_pdf_content = b"%PDF-1.4 encrypted content"
    with patch("app.features.document_parsing.extractors.pdf._parse_pdf_fitz", side_effect=ImportError):
        with patch("app.features.document_parsing.extractors.pdf._parse_pdf_pdfplumber", side_effect=ImportError):
            with patch("pypdf.PdfReader") as mock_reader_cls:
                instance = mock_reader_cls.return_value
                instance.is_encrypted = True
                with pytest.raises(ApiError) as exc_info:
                    parse_pdf_to_blocks(mock_pdf_content)
                assert exc_info.value.status_code == 400
                assert exc_info.value.code == "encrypted_pdf"


def test_pdf_extractor_fitz_encrypted():
    from unittest.mock import MagicMock
    mock_fitz = MagicMock()
    mock_fitz.open.return_value.is_encrypted = True
    with patch.dict("sys.modules", {"fitz": mock_fitz}):
        with pytest.raises(ApiError) as exc_info:
            parse_pdf_to_blocks(b"%PDF-1.4 encrypted")
        assert exc_info.value.status_code == 400
        assert exc_info.value.code == "encrypted_pdf"



# --- 5. DOCX Extractor Tests ---

def test_docx_extractor_paragraphs_and_tables():
    doc = Document()
    doc.add_heading("WORK EXPERIENCE", level=1)
    doc.add_paragraph("Senior Developer at Acme Corp (2020 - Present)")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Project"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Cloud Migration"
    table.cell(1, 1).text = "Lead Architect"

    doc.add_paragraph("Developed core backend microservices.")

    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()

    blocks = parse_docx_to_blocks(docx_bytes)
    assert len(blocks) == 5

    assert blocks[0].block_id == "page-1-block-01"
    assert blocks[0].block_type == "heading"
    assert blocks[0].text == "WORK EXPERIENCE"

    assert blocks[1].block_id == "page-1-block-02"
    assert blocks[1].text == "Senior Developer at Acme Corp (2020 - Present)"
    assert blocks[1].heading_context == "WORK EXPERIENCE"

    assert blocks[2].block_id == "page-1-block-03"
    assert blocks[2].block_type == "table_row"
    assert "[Table Row] Project | Role" in blocks[2].text

    assert blocks[3].block_id == "page-1-block-04"
    assert blocks[3].block_type == "table_row"
    assert "[Table Row] Cloud Migration | Lead Architect" in blocks[3].text

    assert blocks[4].block_id == "page-1-block-05"
    assert blocks[4].text == "Developed core backend microservices."


def test_docx_extractor_corrupted_file():
    invalid_docx = b"Not a valid zip or docx archive"
    with pytest.raises(ApiError) as exc_info:
        parse_docx_to_blocks(invalid_docx)
    assert exc_info.value.status_code == 415
    assert exc_info.value.code == "invalid_docx_archive"


# --- 6. OCR Extractor & Fallback Tests ---

def test_ocr_fallback_unsupported_state():
    with patch("app.features.document_parsing.extractors.ocr.is_ocr_available", return_value=(False, "Tesseract missing")):
        res = process_scanned_pdf(b"%PDF scanned content")
        assert res.status == "OCR_REQUIRED_UNSUPPORTED"
        assert res.is_scanned is True
        assert res.blocks == []
        assert "Optical Character Recognition (OCR) is not installed" in res.message


# --- 7. Unified Extractor Facade Tests ---

def test_unified_extractor_facade_empty_document():
    with pytest.raises(ApiError) as exc_info:
        extract_document_blocks(b"", "resume.pdf", "application/pdf")
    assert exc_info.value.status_code == 400
    assert exc_info.value.code == "empty_document"


def test_unified_extractor_facade_unsupported_type():
    with pytest.raises(ApiError) as exc_info:
        extract_document_blocks(b"data", "resume.txt", "text/plain")
    assert exc_info.value.status_code == 415
    assert exc_info.value.code == "unsupported_document_type"


def test_unified_extractor_facade_docx_success():
    doc = Document()
    doc.add_paragraph("Jane Doe Resume")
    stream = io.BytesIO()
    doc.save(stream)

    res = extract_document_blocks(stream.getvalue(), "jane_resume.docx")
    assert res.status == "SUCCESS"
    assert res.is_scanned is False
    assert len(res.blocks) == 1
    assert res.blocks[0].text == "Jane Doe Resume"
