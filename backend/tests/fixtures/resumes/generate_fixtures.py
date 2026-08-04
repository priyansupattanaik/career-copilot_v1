"""
Synthetic Resume Fixture Generator.

Generates 22 requirement-driven resume fixture files (PDF, DOCX, Scanned, Poor OCR, Empty, Corrupted, Encrypted)
and 22 corresponding golden ground-truth JSON output benchmark files in backend/tests/fixtures/resumes/golden/.
"""

import argparse
import json
import logging
import os
import pathlib
import sys
from dataclasses import dataclass
from typing import Any, Callable, Dict, List, Optional

# Add backend root directory to sys.path if not present
backend_root = str(pathlib.Path(__file__).resolve().parent.parent.parent.parent)
if backend_root not in sys.path:
    sys.path.insert(0, backend_root)

from app.features.document_parsing.schemas import ParsedResumeSchema

# Standard libraries & third-party packages
import pypdf
from PIL import Image, ImageDraw, ImageFilter
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import reportlab
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    FrameBreak,
    HRFlowable,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

FIXTURES_DIR = pathlib.Path(__file__).resolve().parent
GOLDEN_DIR = FIXTURES_DIR / "golden"

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger("generate_fixtures")


class NumberedCanvas(canvas.Canvas):
    """Two-pass canvas for adding running footers like 'Page X of Y'."""

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        super().__init__(*args, **kwargs)
        self._saved_page_states: List[Dict[str, Any]] = []

    def showPage(self) -> None:
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self) -> None:
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count: int) -> None:
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#666666"))
        self.drawRightString(letter[0] - 36, 36, f"Page {self._pageNumber} of {page_count}")
        self.restoreState()


@dataclass
class FixtureSpec:
    id: int
    slug: str
    filename: str
    golden_filename: str
    file_type: str  # "pdf" | "docx"
    description: str
    generator_fn: Callable[[pathlib.Path, Dict[str, Any]], None]
    golden_data: Dict[str, Any]

    @property
    def fixture_path(self) -> pathlib.Path:
        return FIXTURES_DIR / self.filename

    @property
    def golden_path(self) -> pathlib.Path:
        return GOLDEN_DIR / self.golden_filename


def make_field(value: Any, block_ids: List[str], confidence: str = "HIGH", warning: Optional[str] = None) -> Dict[str, Any]:
    return {
        "value": value,
        "evidence_block_ids": block_ids,
        "confidence": confidence,
        "warning": warning,
    }


def make_empty_field(warning: Optional[str] = None) -> Dict[str, Any]:
    """Helper to generate a valid empty FieldWrapper dict for absent fields."""
    return {
        "value": None,
        "evidence_block_ids": [],
        "confidence": "HIGH",
        "warning": warning,
    }


def normalize_golden_benchmark(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Normalizes golden benchmark dictionaries to align with ParsedResumeSchema field names
    and replaces raw null values with empty FieldWrapper dictionaries.
    """
    if not isinstance(data, dict):
        return data

    if "value" in data and "evidence_block_ids" in data:
        return data

    normalized = {}
    for key, val in data.items():
        # Correct field name mismatches
        norm_key = key
        if key == "licenses":
            norm_key = "licences"
        elif key == "current_role":
            norm_key = "is_current"
        elif key == "field":
            norm_key = "field_of_study"
        elif key == "project_name":
            norm_key = "name"

        # Filter out misplaced top-level project keys
        if key in ("dates", "confidence", "evidence_block_ids") and not isinstance(val, (dict, list)):
            continue

        if val is None:
            normalized[norm_key] = make_empty_field()
        elif isinstance(val, dict):
            normalized[norm_key] = normalize_golden_benchmark(val)
        elif isinstance(val, list):
            normalized[norm_key] = [normalize_golden_benchmark(item) for item in val]
        else:
            normalized[norm_key] = val

    return normalized


# ==============================================================================
# Generator Functions for 22 Fixtures
# ==============================================================================

# --- Fixture 01: Single Column PDF ---
def gen_01_single_column(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('DocTitle', parent=styles['Title'], fontSize=20, leading=24, textColor=colors.HexColor('#1A252C'), alignment=0)
    heading_style = ParagraphStyle('SectionHeader', parent=styles['Heading2'], fontSize=12, leading=16, textColor=colors.HexColor('#007AFF'), spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14, textColor=colors.HexColor('#333333'))

    story = [
        Paragraph("Alex Mercer", title_style),
        Paragraph("Email: alex.mercer@email.com | Phone: (555) 019-2834 | San Francisco, CA", body_style),
        Paragraph("LinkedIn: linkedin.com/in/alexmercer | GitHub: github.com/alexmercer", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC'), spaceBefore=6, spaceAfter=10),
        Paragraph("PROFESSIONAL SUMMARY", heading_style),
        Paragraph("Senior Backend Engineer with 8+ years of experience building scalable microservices in Python and Go.", body_style),
        Spacer(1, 10),
        Paragraph("SKILLS", heading_style),
        Paragraph("Python, FastAPI, Django, Go, PostgreSQL, Redis, Docker, Kubernetes, AWS", body_style),
        Spacer(1, 10),
        Paragraph("WORK EXPERIENCE", heading_style),
        Paragraph("<b>Senior Backend Engineer</b> — TechCorp (2021 – Present)", body_style),
        Paragraph("• Architected high-throughput REST APIs handling 50M daily requests.", body_style),
        Paragraph("• Optimized database queries, reducing API latency by 35%.", body_style),
        Spacer(1, 6),
        Paragraph("<b>Software Engineer</b> — DataFlow Inc (2018 – 2021)", body_style),
        Paragraph("• Built real-time telemetry streaming pipelines using Kafka and Python.", body_style),
        Spacer(1, 10),
        Paragraph("EDUCATION", heading_style),
        Paragraph("<b>B.S. in Computer Science</b> — University of California, Berkeley (2014 – 2018)", body_style),
    ]
    doc.build(story)


# --- Fixture 02: Two Column PDF ---
def gen_02_two_column(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = BaseDocTemplate(str(path), pagesize=letter)
    margin = 36
    page_w, page_h = letter
    col_w = (page_w - 2 * margin - 18) / 2

    frame1 = Frame(margin, margin, col_w, page_h - 2 * margin, id='col1', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    frame2 = Frame(margin + col_w + 18, margin, col_w, page_h - 2 * margin, id='col2', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    template = PageTemplate(id='two_col', frames=[frame1, frame2])
    doc.addPageTemplates([template])

    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle('ColH1', parent=styles['Heading1'], fontSize=16, leading=20, textColor=colors.HexColor('#1F4E79'))
    h2_style = ParagraphStyle('ColH2', parent=styles['Heading2'], fontSize=12, leading=16, textColor=colors.HexColor('#1F4E79'), spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle('ColBody', parent=styles['Normal'], fontSize=9.5, leading=13, textColor=colors.HexColor('#333333'))

    story = [
        Paragraph("Elena Rostova", h1_style),
        Paragraph("Lead Data Scientist | elena.rostova@example.com | Boston, MA", body_style),
        Spacer(1, 10),
        Paragraph("WORK EXPERIENCE", h2_style),
        Paragraph("<b>Lead Data Scientist</b> — BioTech AI Labs (2020 – Present)", body_style),
        Paragraph("• Developed deep learning models for genomic sequence analysis.", body_style),
        Paragraph("• Deployed PyTorch models on AWS SageMaker with 99.9% uptime.", body_style),
        Spacer(1, 8),
        Paragraph("<b>Data Scientist</b> — HealthData Systems (2017 – 2020)", body_style),
        Paragraph("• Built predictive models for patient readmission risks using XGBoost.", body_style),
        Spacer(1, 10),
        Paragraph("EDUCATION", h2_style),
        Paragraph("<b>Ph.D. in Computational Biology</b> — MIT (2012 – 2017)", body_style),
        FrameBreak(),
        Paragraph("SKILLS & TOOLKIT", h2_style),
        Paragraph("<b>Languages:</b> Python, R, SQL, C++", body_style),
        Paragraph("<b>ML Frameworks:</b> PyTorch, TensorFlow, Scikit-Learn", body_style),
        Paragraph("<b>Cloud & Ops:</b> AWS SageMaker, Docker, MLflow", body_style),
        Spacer(1, 10),
        Paragraph("CERTIFICATIONS", h2_style),
        Paragraph("• AWS Certified Machine Learning - Specialty (2022)", body_style),
        Spacer(1, 10),
        Paragraph("PUBLICATIONS", h2_style),
        Paragraph("• Rostova E., et al. 'Genomic Sequence Transformer Networks', Nature AI, 2021.", body_style),
    ]
    doc.build(story)


# --- Fixture 03: Sidebar PDF ---
def draw_sidebar_bg(canvas_obj: canvas.Canvas, doc_obj: Any) -> None:
    canvas_obj.saveState()
    canvas_obj.setFillColor(colors.HexColor('#2C3E50'))
    canvas_obj.rect(0, 0, 180, letter[1], fill=1, stroke=0)
    canvas_obj.restoreState()


def gen_03_sidebar(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = BaseDocTemplate(str(path), pagesize=letter)
    sidebar_frame = Frame(18, 18, 144, letter[1] - 36, id='sidebar', leftPadding=6, rightPadding=6, topPadding=18)
    main_frame = Frame(198, 18, letter[0] - 216, letter[1] - 36, id='main', leftPadding=6, rightPadding=6, topPadding=18)

    template = PageTemplate(id='sidebar_layout', frames=[sidebar_frame, main_frame], onPage=draw_sidebar_bg)
    doc.addPageTemplates([template])

    styles = getSampleStyleSheet()
    side_title = ParagraphStyle('SideTitle', parent=styles['Heading2'], fontSize=11, leading=14, textColor=colors.white, spaceBefore=8, spaceAfter=4)
    side_body = ParagraphStyle('SideBody', parent=styles['Normal'], fontSize=9, leading=13, textColor=colors.HexColor('#ECEFF1'))
    main_title = ParagraphStyle('MainTitle', parent=styles['Title'], fontSize=20, leading=24, textColor=colors.HexColor('#2C3E50'), alignment=0)
    main_h2 = ParagraphStyle('MainH2', parent=styles['Heading2'], fontSize=12, leading=16, textColor=colors.HexColor('#2C3E50'), spaceBefore=10, spaceAfter=4)
    main_body = ParagraphStyle('MainBody', parent=styles['Normal'], fontSize=9.5, leading=13.5, textColor=colors.HexColor('#333333'))

    story = [
        Paragraph("CONTACT", side_title),
        Paragraph("marcus.vance@example.com", side_body),
        Paragraph("+1 (555) 019-4820", side_body),
        Paragraph("Austin, TX", side_body),
        Paragraph("linkedin.com/in/mvance", side_body),
        Spacer(1, 12),
        Paragraph("SKILLS", side_title),
        Paragraph("• Python & Go", side_body),
        Paragraph("• FastAPI & Django", side_body),
        Paragraph("• AWS & Terraform", side_body),
        Paragraph("• PostgreSQL & Redis", side_body),
        Spacer(1, 12),
        Paragraph("LANGUAGES", side_title),
        Paragraph("English (Native)", side_body),
        Paragraph("Spanish (Fluent)", side_body),
        FrameBreak(),
        Paragraph("Marcus Vance", main_title),
        Paragraph("Principal Cloud Infrastructure Systems Engineer", ParagraphStyle('Sub', parent=main_body, textColor=colors.HexColor('#7F8C8D'), fontSize=11)),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#BDC3C7'), spaceBefore=4, spaceAfter=10),
        Paragraph("SUMMARY", main_h2),
        Paragraph("Principal Cloud Infrastructure Systems Engineer with over 10 years experience designing high-availability distributed systems.", main_body),
        Spacer(1, 8),
        Paragraph("EXPERIENCE", main_h2),
        Paragraph("<b>Principal Cloud Engineer</b> — CloudScale Systems (2020 – Present)", main_body),
        Paragraph("• Designed Kubernetes orchestration platform automating multiregion deployments.", main_body),
        Paragraph("• Reduced infrastructure AWS costs by 28% through reserved instance planning.", main_body),
        Spacer(1, 6),
        Paragraph("<b>Senior DevOps Engineer</b> — DataCore Corp (2016 – 2020)", main_body),
        Paragraph("• Built CI/CD automation pipelines serving 120+ microservices.", main_body),
    ]
    doc.build(story)


# --- Fixture 04: Table Based PDF ---
def gen_04_table_based(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TTitle', parent=styles['Heading1'], fontSize=16, leading=20, textColor=colors.HexColor('#1A252C'))
    cell_bold = ParagraphStyle('CBold', parent=styles['Normal'], fontSize=9.5, leading=13, fontName='Helvetica-Bold')
    cell_body = ParagraphStyle('CBody', parent=styles['Normal'], fontSize=9.5, leading=13)

    data = [
        [Paragraph("<b>CANDIDATE PROFILE</b>", title_style), Paragraph("<b>Contact:</b> samuel.oak@example.com | (555) 321-9876", cell_body)],
        [Paragraph("SUMMARY", cell_bold), Paragraph("Experienced Full Stack Developer specializing in React, Node.js, and Python web architectures.", cell_body)],
        [Paragraph("SKILLS", cell_bold), Paragraph("JavaScript, TypeScript, React, Node.js, Python, PostgreSQL, GraphQL, Docker", cell_body)],
        [Paragraph("EXPERIENCE", cell_bold), Paragraph("<b>Full Stack Lead</b> — WebMatrix Inc (2019 - Present)<br/>• Led development of SaaS enterprise dashboard.<br/>• Integrated Stripe payments and OAuth authentication.", cell_body)],
        [Paragraph("EDUCATION", cell_bold), Paragraph("<b>B.S. Software Engineering</b> — Oregon State University (2015 - 2019)", cell_body)],
    ]
    t = Table(data, colWidths=[120, 420])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F2F4F7')),
        ('GRID', (0, 0), (-1, -1), 0.75, colors.HexColor('#D0D5DD')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    doc.build([t])


# --- Fixture 05: Long Multipage PDF ---
def gen_05_long_multipage(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle('H1', parent=styles['Title'], fontSize=20, leading=24, textColor=colors.HexColor('#1A252C'), alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=12, leading=16, textColor=colors.HexColor('#007AFF'), spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14)

    story = [
        Paragraph("Dr. Arthur Pendelton", h1_style),
        Paragraph("Vice President of Engineering | arthur.pendelton@executive.org | (555) 999-0001 | New York, NY", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC'), spaceAfter=10),
        Paragraph("EXECUTIVE SUMMARY", h2_style),
        Paragraph("Global Engineering Executive with 20+ years of leadership scaling engineering organizations from 20 to 500+ engineers across Americas, EMEA, and APAC.", body_style),
        Spacer(1, 10),
        Paragraph("CAREER EXPERIENCE (PAGE 1)", h2_style),
        Paragraph("<b>VP of Engineering</b> — Global Tech Enterprises (2020 – Present)", body_style),
        Paragraph("• Oversaw $80M annual engineering budget and managed 450+ engineering headcount.", body_style),
        Paragraph("• Spearheaded cloud transformation migrating legacy on-prem systems to AWS.", body_style),
        Spacer(1, 15),
        Paragraph("<b>Senior Director of Software Architecture</b> — Enterprise Cloud Corp (2015 – 2020)", body_style),
        Paragraph("• Directed software architecture strategy across 12 product lines.", body_style),
        PageBreak(),  # Page 2
        Paragraph("CAREER EXPERIENCE CONTINUED (PAGE 2)", h2_style),
        Paragraph("<b>Director of Engineering</b> — NextGen Systems (2010 – 2015)", body_style),
        Paragraph("• Led development of core microservices platform processing $2B in transaction volume.", body_style),
        Spacer(1, 10),
        Paragraph("<b>Principal Software Architect</b> — Legacy Technologies (2005 – 2010)", body_style),
        Paragraph("• Designed high-performance C++ real-time order matching engine.", body_style),
        Spacer(1, 15),
        Paragraph("BOARD ADVISORY & FELLOWSHIPS", h2_style),
        Paragraph("• Technical Advisory Board Member — Cloud Native Computing Foundation (CNCF)", body_style),
        Paragraph("• Executive Fellow — Institute of Software Engineering", body_style),
        PageBreak(),  # Page 3
        Paragraph("EDUCATION & PATENTS (PAGE 3)", h2_style),
        Paragraph("<b>Ph.D. in Computer Engineering</b> — Stanford University (2001 – 2005)", body_style),
        Paragraph("<b>M.S. in Electrical Engineering</b> — MIT (1999 – 2001)", body_style),
        Spacer(1, 10),
        Paragraph("ISSUED PATENTS", h2_style),
        Paragraph("• US Patent 9,842,102: Distributed Consensus protocol for high-latency networks.", body_style),
        Paragraph("• US Patent 8,912,441: Dynamic load-balancing across edge nodes.", body_style),
    ]
    doc.build(story, canvasmaker=NumberedCanvas)


# --- Fixture 06: Minimal Fresher PDF ---
def gen_06_minimal_fresher(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#007AFF'), spaceBefore=8, spaceAfter=2)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("Sam Rivera", title_style),
        Paragraph("Email: sam.rivera@gmail.com | Phone: (555) 444-3322 | Chicago, IL", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC'), spaceAfter=8),
        Paragraph("OBJECTIVE", h2_style),
        Paragraph("Enthusiastic Computer Science graduate seeking an entry-level Junior Software Developer position.", body_style),
        Spacer(1, 8),
        Paragraph("EDUCATION", h2_style),
        Paragraph("<b>B.S. in Computer Science</b> — University of Illinois Chicago (Graduated May 2024)", body_style),
        Paragraph("GPA: 3.7 / 4.0", body_style),
        Spacer(1, 8),
        Paragraph("TECHNICAL SKILLS", h2_style),
        Paragraph("Python, Java, HTML/CSS, JavaScript, Git, SQL", body_style),
        Spacer(1, 8),
        Paragraph("ACADEMIC PROJECTS", h2_style),
        Paragraph("<b>Student Portal App</b> (Senior Capstone Project)", body_style),
        Paragraph("• Created a web portal using Python Flask and SQLite for managing course assignments.", body_style),
    ]
    doc.build(story)


# --- Fixture 07: Senior Technical PDF ---
def gen_07_senior_technical(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=20, leading=24, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=12, leading=16, textColor=colors.HexColor('#107C41'), spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13.5)

    story = [
        Paragraph("Vikram Patel", title_style),
        Paragraph("Staff Distributed Systems Architect | vikram.patel@techleader.io | Seattle, WA | github.com/vpatel-staff", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#107C41'), spaceAfter=8),
        Paragraph("SUMMARY", h2_style),
        Paragraph("Staff Distributed Systems Architect with 12+ years of experience designing fault-tolerant cloud infrastructures, event-driven architectures, and high-throughput streaming platforms.", body_style),
        Spacer(1, 8),
        Paragraph("TECHNICAL EXPERTISE", h2_style),
        Paragraph("<b>Languages:</b> Go, Rust, Python, C++, SQL, ProtoBuf", body_style),
        Paragraph("<b>Distributed Systems:</b> Apache Kafka, gRPC, Kubernetes, Istio Service Mesh, Redis, Cassandra", body_style),
        Paragraph("<b>Cloud Infrastructure:</b> AWS (EKS, DynamoDB, CloudFront), Terraform, Prometheus, Grafana", body_style),
        Spacer(1, 8),
        Paragraph("PROFESSIONAL EXPERIENCE", h2_style),
        Paragraph("<b>Staff Engineer & Architect</b> — HyperScale Systems (2021 – Present)", body_style),
        Paragraph("• Lead architect for real-time payment settlement engine processing 100,000 TPS with 99.999% availability.", body_style),
        Paragraph("• Authored company-wide technical RFCs for zero-trust microservice mesh architecture.", body_style),
        Spacer(1, 6),
        Paragraph("<b>Principal Backend Engineer</b> — DataStream Inc (2017 – 2021)", body_style),
        Paragraph("• Engineered distributed event stream processing platform using Go and Apache Kafka.", body_style),
        Spacer(1, 6),
        Paragraph("<b>Senior Systems Engineer</b> — CloudNetworks (2014 – 2017)", body_style),
        Paragraph("• Implemented custom eBPF network packet filter reducing DDoS attack impact by 90%.", body_style),
        Spacer(1, 8),
        Paragraph("EDUCATION & CERTIFICATIONS", h2_style),
        Paragraph("<b>M.S. in Computer Science</b> — University of Washington (2012 – 2014)", body_style),
        Paragraph("Certified Kubernetes Administrator (CKA)", body_style),
    ]
    doc.build(story)


# --- Fixture 08: Career Change PDF ---
def gen_08_career_change(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#D9381E'), spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("Jordan Lee", title_style),
        Paragraph("Full Stack Developer (Transitioned from Financial Analysis) | jordan.lee@email.com | New York, NY", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#D9381E'), spaceAfter=8),
        Paragraph("CAREER OBJECTIVE & PROFILE", h2_style),
        Paragraph("Former Financial Analyst transitioning into Software Engineering. Combines strong quantitative analytical skills with rigorous full-stack web development expertise.", body_style),
        Spacer(1, 8),
        Paragraph("TRANSFERABLE & TECHNICAL SKILLS", h2_style),
        Paragraph("<b>Software Engineering:</b> Python, JavaScript, React, Node.js, Express, PostgreSQL, Git", body_style),
        Paragraph("<b>Financial & Data Analysis:</b> SQL, Financial Modeling, Statistical Analysis, Excel VBA", body_style),
        Spacer(1, 8),
        Paragraph("TECHNICAL PROJECTS", h2_style),
        Paragraph("<b>Portfolio Analytics Dashboard</b> (Full Stack Web Application)", body_style),
        Paragraph("• Built React frontend with Recharts visualization and Python FastAPI backend for stock portfolio tracking.", body_style),
        Spacer(1, 8),
        Paragraph("PREVIOUS PROFESSIONAL EXPERIENCE", h2_style),
        Paragraph("<b>Senior Financial Analyst</b> — WallStreet Capital (2018 – 2023)", body_style),
        Paragraph("• Built automated Python financial reporting scripts, cutting monthly close time by 40 hours.", body_style),
        Spacer(1, 8),
        Paragraph("EDUCATION & IMMERSIVE TRAINING", h2_style),
        Paragraph("<b>Full-Stack Software Engineering Certificate</b> — General Assembly (2023)", body_style),
        Paragraph("<b>B.S. in Finance</b> — NYU Stern School of Business (2014 – 2018)", body_style),
    ]
    doc.build(story)


# --- Fixture 09: Academic CV PDF ---
def gen_09_academic_cv(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=20, leading=24, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=12, leading=16, textColor=colors.HexColor('#4A154B'), spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13.5)

    story = [
        Paragraph("Dr. Catherine Thorne", title_style),
        Paragraph("Associate Professor of Computer Science | catherine.thorne@university.edu | Cambridge, MA", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#4A154B'), spaceAfter=8),
        Paragraph("RESEARCH INTERESTS", h2_style),
        Paragraph("Natural Language Processing, Program Synthesis, Formal Verification of Neural Networks.", body_style),
        Spacer(1, 8),
        Paragraph("EDUCATION", h2_style),
        Paragraph("<b>Ph.D. in Computer Science</b> — Harvard University (2015)", body_style),
        Paragraph("Dissertation: 'Grounded Semantic Parsing for Programming Languages'", body_style),
        Paragraph("<b>B.S. in Mathematics & CS</b> — Yale University (2010)", body_style),
        Spacer(1, 8),
        Paragraph("ACADEMIC APPOINTMENTS", h2_style),
        Paragraph("<b>Associate Professor</b> — Dept of CS, Harvard University (2021 – Present)", body_style),
        Paragraph("<b>Assistant Professor</b> — Dept of CS, Harvard University (2015 – 2021)", body_style),
        PageBreak(),  # Page 2
        Paragraph("PUBLICATIONS & CONFERENCES", h2_style),
        Paragraph("• Thorne C., & Vaswani A. 'Neural Program Synthesis with Constraint Verification', NeurIPS 2022.", body_style),
        Paragraph("• Thorne C. 'Semantic Grounding in Large Code Models', ACL 2020.", body_style),
        Spacer(1, 8),
        Paragraph("RESEARCH GRANTS & AWARDS", h2_style),
        Paragraph("• NSF CAREER Award: Automated Code Verification via Grounded Models ($550,000) — 2020", body_style),
        Spacer(1, 8),
        Paragraph("TEACHING EXPERIENCE", h2_style),
        Paragraph("• CS 281: Advanced Natural Language Processing (Graduate Level)", body_style),
        Paragraph("• CS 182: Artificial Intelligence Principles", body_style),
    ]
    doc.build(story, canvasmaker=NumberedCanvas)


# --- Fixture 10: Project Heavy PDF ---
def gen_10_project_heavy(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#008080'), spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("David Kim", title_style),
        Paragraph("Open Source Contributor & Software Developer | david.kim@dev.io | github.com/dkim-dev", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#008080'), spaceAfter=8),
        Paragraph("SKILLS OVERVIEW", h2_style),
        Paragraph("Python, Rust, TypeScript, React, Docker, WebAssembly, GraphQL, PostgreSQL", body_style),
        Spacer(1, 8),
        Paragraph("FEATURED OPEN SOURCE PROJECTS", h2_style),
        Paragraph("<b>1. FastGraphDB</b> (High Performance Graph Database Engine)", body_style),
        Paragraph("• Description: Rust-based graph storage engine supporting Cypher queries.", body_style),
        Paragraph("• Tech Stack: Rust, Tokio, RocksDB | Stars: 2.4k GitHub Stars", body_style),
        Spacer(1, 6),
        Paragraph("<b>2. PyCleanCode</b> (Static Analysis Linter for Python)", body_style),
        Paragraph("• Description: Fast Python linter enforcing PEP8 and grounding verification rules.", body_style),
        Paragraph("• Tech Stack: Python, AST Parsing | Downloads: 50,000+ PyPI downloads", body_style),
        Spacer(1, 6),
        Paragraph("<b>3. ReactCanvasFlow</b> (UI Flowchart Diagramming Library)", body_style),
        Paragraph("• Description: Canvas-rendered UI component library for complex workflow graphs.", body_style),
        Paragraph("• Tech Stack: TypeScript, React, HTML5 Canvas", body_style),
        Spacer(1, 6),
        Paragraph("<b>4. MicroRunner</b> (Lightweight Docker Container Runner)", body_style),
        Paragraph("• Description: Minimal container orchestration daemon for edge devices.", body_style),
        Paragraph("• Tech Stack: Go, Docker Engine API", body_style),
        Spacer(1, 8),
        Paragraph("EDUCATION", h2_style),
        Paragraph("<b>B.S. in Computer Engineering</b> — Georgia Tech (2018 – 2022)", body_style),
    ]
    doc.build(story)


# --- Fixture 11: Freelance PDF ---
def gen_11_freelance(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#8E44AD'), spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("Maya Lin", title_style),
        Paragraph("Independent Cloud Solutions Architect & Consultant | maya@linconsulting.com | San Jose, CA", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#8E44AD'), spaceAfter=8),
        Paragraph("CONSULTING SERVICES", h2_style),
        Paragraph("Cloud Infrastructure Migration, AWS Cost Optimization, DevOps Automation, Serverless Architecture.", body_style),
        Spacer(1, 8),
        Paragraph("INDEPENDENT CLIENT ENGAGEMENTS", h2_style),
        Paragraph("<b>FinTech Client Alpha</b> — Senior Cloud Consultant (Contract, 2023)", body_style),
        Paragraph("• Migrated monolithic payment backend to AWS Lambda and DynamoDB serverless setup.", body_style),
        Spacer(1, 6),
        Paragraph("<b>HealthTech Client Beta</b> — DevOps Specialist (Contract, 2022 – 2023)", body_style),
        Paragraph("• Implemented HIPAA-compliant Kubernetes cluster on GCP using Terraform.", body_style),
        Spacer(1, 6),
        Paragraph("<b>E-Commerce Client Gamma</b> — Infrastructure Architect (Contract, 2021 – 2022)", body_style),
        Paragraph("• Re-architected database caching using Redis, resolving Black Friday traffic bottlenecks.", body_style),
        Spacer(1, 8),
        Paragraph("TECHNICAL SKILLS", h2_style),
        Paragraph("AWS, GCP, Terraform, Kubernetes, Docker, Python, Node.js, CI/CD pipelines", body_style),
    ]
    doc.build(story)


# --- Fixture 12: Multiple Roles PDF ---
def gen_12_multiple_roles(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#1A252C'), spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("Chris Evans", title_style),
        Paragraph("Staff Software Engineer | chris.evans@email.com | Austin, TX", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC'), spaceAfter=8),
        Paragraph("WORK EXPERIENCE", h2_style),
        Paragraph("<b>TechCorp Solutions</b> (Total Tenure: 2018 – Present)", body_style),
        Paragraph("<i>Staff Software Engineer</i> (2022 – Present)", body_style),
        Paragraph("• Promoted to Staff Engineer to drive technical roadmap for enterprise core backend.", body_style),
        Paragraph("• Mentored 15+ software engineers across 3 frontend/backend squads.", body_style),
        Spacer(1, 4),
        Paragraph("<i>Senior Software Engineer</i> (2020 – 2022)", body_style),
        Paragraph("• Promoted from Engineer II after leading successful migration to Python 3.10 and FastAPI.", body_style),
        Spacer(1, 4),
        Paragraph("<i>Software Engineer II</i> (2018 – 2020)", body_style),
        Paragraph("• Developed core REST APIs for billing and subscription management.", body_style),
        Spacer(1, 8),
        Paragraph("EDUCATION", h2_style),
        Paragraph("<b>B.S. in Computer Science</b> — University of Texas at Austin (2014 – 2018)", body_style),
    ]
    doc.build(story)


# --- Fixture 13: Overlapping Dates PDF ---
def gen_13_overlapping_dates(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#1A252C'), spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("Taylor Morgan", title_style),
        Paragraph("Full Stack Engineer & Tech Advisor | taylor.morgan@email.com | Denver, CO", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC'), spaceAfter=8),
        Paragraph("EXPERIENCE & CONCURRENT ROLES", h2_style),
        Paragraph("<b>Full Stack Lead Engineer</b> — Apex Innovations (Jan 2021 – Present)", body_style),
        Paragraph("• Leading full stack development team building SaaS platform.", body_style),
        Spacer(1, 6),
        Paragraph("<b>Technical Co-Founder & Advisor</b> — Startup Accelerator Project (Mar 2021 – Dec 2022)", body_style),
        Paragraph("• Concurrent co-founder role building MVP web app for early-stage stealth startup.", body_style),
        Spacer(1, 6),
        Paragraph("<b>Part-time Consulting Engineer</b> — Cloud Solutions LLC (Jun 2021 – Jan 2022)", body_style),
        Paragraph("• Concurrent evening consulting role auditing AWS IAM security policies.", body_style),
    ]
    doc.build(story)


# --- Fixture 14: Unusual Headings PDF ---
def gen_14_unusual_headings(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#C0392B'), spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("Morgan Reed", title_style),
        Paragraph("Creative Developer & Designer | morgan.reed@creative.net", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#C0392B'), spaceAfter=8),
        Paragraph("WHERE I HAVE WORKED", h2_style),
        Paragraph("<b>Lead Interactive Developer</b> — PixelCraft Studios (2020 – Present)", body_style),
        Paragraph("• Designed interactive 3D WebGL web experiences.", body_style),
        Spacer(1, 8),
        Paragraph("MY TOOLKIT & CAPABILITIES", h2_style),
        Paragraph("JavaScript, Three.js, React, WebGL, Python, CSS Canvas, HTML5", body_style),
        Spacer(1, 8),
        Paragraph("SCHOLASTIC BACKGROUND", h2_style),
        Paragraph("<b>B.F.A. in Digital Media Arts</b> — Rhode Island School of Design (2016 – 2020)", body_style),
        Spacer(1, 8),
        Paragraph("STUFF I BUILT", h2_style),
        Paragraph("<b>Interactive Shader Generator</b> — Open source WebGL shader tool with 1,000+ stars.", body_style),
    ]
    doc.build(story)


# --- Fixture 15: No Headings PDF ---
def gen_15_no_headings(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14)

    story = [
        Paragraph("Riley Harper", title_style),
        Paragraph("Software Developer | riley.harper@email.com | (555) 777-8899 | San Diego, CA", body_style),
        Spacer(1, 10),
        Paragraph("Software engineer with 5 years experience specializing in web APIs and frontend user interfaces.", body_style),
        Spacer(1, 10),
        Paragraph("Software Engineer at WebWorks Corp from 2021 to Present. Developed Python FastAPI services and React components. Improved web app loading performance by 40%.", body_style),
        Spacer(1, 10),
        Paragraph("Junior Developer at CodeLab from 2019 to 2021. Built client websites using HTML, CSS, JavaScript, and Node.js.", body_style),
        Spacer(1, 10),
        Paragraph("Proficient in Python, JavaScript, TypeScript, React, HTML, CSS, SQL, Git, and Docker.", body_style),
        Spacer(1, 10),
        Paragraph("Bachelor of Science in Computer Science from San Diego State University, 2015 to 2019.", body_style),
    ]
    doc.build(story)


# --- Fixture 16: Icons PDF ---
def gen_16_icons(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, leading=22, alignment=0)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11, leading=15, textColor=colors.HexColor('#27AE60'), spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13)

    story = [
        Paragraph("Avery Bennett", title_style),
        Paragraph("📧 avery.bennett@email.com | 📱 (555) 019-7733 | 🌐 averybennett.dev | 📍 Portland, OR", body_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#27AE60'), spaceAfter=8),
        Paragraph("💼 PROFESSIONAL EXPERIENCE", h2_style),
        Paragraph("<b>Senior Backend Developer</b> — EcoTech Systems (2020 – Present)", body_style),
        Paragraph("• 🛠️ Built REST APIs using Python and Django.", body_style),
        Paragraph("• ⚡ Optimized PostgreSQL database queries.", body_style),
        Spacer(1, 8),
        Paragraph("🎓 EDUCATION", h2_style),
        Paragraph("<b>B.S. Computer Science</b> — University of Oregon (2016 – 2020)", body_style),
    ]
    doc.build(story)


# --- Fixture 17: DOCX Tables ---
def gen_17_docx_tables(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    p_title = doc.add_paragraph()
    r_name = p_title.add_run("Marcus Vance")
    r_name.font.size = Pt(20)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(31, 78, 121)

    p_contact = doc.add_paragraph("Email: marcus.vance@example.com | Phone: (555) 019-4820 | Location: Austin, TX")
    p_contact.paragraph_format.space_after = Pt(12)

    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    tbl_sum = doc.add_table(rows=1, cols=1)
    tbl_sum.style = 'Table Grid'
    c_sum = tbl_sum.rows[0].cells[0]
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="F2F4F7"/>'
    c_sum._tc.get_or_add_tcPr().append(parse_xml(shd_xml))
    c_sum.paragraphs[0].text = "Principal Systems Engineer with over 10 years experience designing distributed systems, cloud architecture, and microservices."

    doc.add_heading("EMPLOYMENT MATRIX", level=1)
    tbl_exp = doc.add_table(rows=1, cols=4)
    tbl_exp.style = 'Table Grid'
    hdr = tbl_exp.rows[0].cells
    hdr_titles = ["Employer / Dates", "Role Title", "Responsibilities & Achievements", "Technologies"]
    for idx, t_text in enumerate(hdr_titles):
        hdr[idx].text = t_text
        hdr[idx]._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>'))
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    row1 = tbl_exp.add_row().cells
    row1[0].text = "CloudScale Inc.\n2021 - Present"
    row1[1].text = "Lead Architect"
    row1[2].text = "• Managed team of 8 backend engineers.\n• Designed Kubernetes platform."
    row1[3].text = "Python, FastAPI, AWS, K8s, Docker"

    doc.save(str(path))


# --- Fixture 18: Scanned PDF ---
def gen_18_scanned(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    img = Image.new('RGB', (1654, 2339), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((100, 100), "SCANNED RESUME - SARAH CONNOR", fill=(0, 0, 0))
    draw.text((100, 160), "Contact: sarah.connor@cyberdyne.org | Phone: 555-0199", fill=(30, 30, 30))
    draw.text((100, 220), "EXPERIENCE", fill=(0, 0, 0))
    draw.text((100, 260), "Cyberdyne Systems - Security Specialist (2020 - 2024)", fill=(30, 30, 30))
    draw.text((100, 300), "• Implemented defensive security protocols and threat monitoring.", fill=(50, 50, 50))
    img.save(str(path), "PDF", resolution=200.0)


# --- Fixture 19: Poor OCR PDF ---
def gen_19_poor_ocr(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    img = Image.new('RGB', (612, 792), color=(240, 240, 235))
    draw = ImageDraw.Draw(img)
    draw.text((50, 50), "P00r QUA1ITY 0CR R3SUM3 - J0HN D0E", fill=(80, 80, 80))
    draw.text((50, 90), "Exp3ri3nc3: S0ftw4r3 Eng1n33r @ T3ch (2021-2023)", fill=(90, 90, 90))
    draw.text((50, 130), "Sk1lls: Pyth0n, J4v4Scr1pt, D0ck3r", fill=(90, 90, 90))
    img = img.filter(ImageFilter.GaussianBlur(radius=1.2))
    img.save(str(path), "PDF", resolution=75.0, quality=15)


# --- Fixture 20: Empty PDF ---
def gen_20_empty(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    with open(path, "wb") as f:
        f.write(b"")


# --- Fixture 21: Corrupted PDF ---
def gen_21_corrupted(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    corrupted_bytes = (
        b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
        b"CORRUPTED_FILE_INVALID_XREF_TRAILER_STREAM_1234567890_BYTES_PAYLOAD"
    )
    with open(path, "wb") as f:
        f.write(corrupted_bytes)


# --- Fixture 22: Encrypted PDF ---
def gen_22_encrypted(path: pathlib.Path, golden: Dict[str, Any]) -> None:
    temp_path = str(path) + ".tmp"
    c = canvas.Canvas(temp_path, pagesize=letter)
    c.drawString(100, 700, "Encrypted Resume Content - Confidential Candidate Data")
    c.save()

    reader = pypdf.PdfReader(temp_path)
    writer = pypdf.PdfWriter()
    writer.append(reader)
    writer.encrypt(user_password="user123", owner_password="owner123")

    with open(path, "wb") as f:
        writer.write(f)

    if os.path.exists(temp_path):
        os.remove(temp_path)


# ==============================================================================
# 22 Golden Ground Truth Definitions
# ==============================================================================

GOLDEN_BENCHMARKS: Dict[int, Dict[str, Any]] = {
    1: {
        "fixture_meta": {"fixture_id": 1, "fixture_slug": "01_single_column", "filename": "01_single_column.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Alex Mercer", ["page-1-block-01"]),
            "email": make_field("alex.mercer@email.com", ["page-1-block-02"]),
            "phone": make_field("(555) 019-2834", ["page-1-block-02"]),
            "location": make_field("San Francisco, CA", ["page-1-block-02"]),
            "linkedin": make_field("linkedin.com/in/alexmercer", ["page-1-block-03"]),
            "github": make_field("github.com/alexmercer", ["page-1-block-03"]),
            "portfolio": None,
            "other_links": [],
        },
        "professional_summary": make_field("Senior Backend Engineer with 8+ years of experience building scalable microservices in Python and Go.", ["page-1-block-05"]),
        "target_role": make_field("Senior Backend Engineer", ["page-1-block-05"]),
        "skills": [
            {"name": make_field("Python", ["page-1-block-07"]), "category": make_field("Languages", ["page-1-block-07"])},
            {"name": make_field("FastAPI", ["page-1-block-07"]), "category": make_field("Frameworks", ["page-1-block-07"])},
            {"name": make_field("Go", ["page-1-block-07"]), "category": make_field("Languages", ["page-1-block-07"])},
            {"name": make_field("PostgreSQL", ["page-1-block-07"]), "category": make_field("Databases", ["page-1-block-07"])},
            {"name": make_field("Docker", ["page-1-block-07"]), "category": make_field("DevOps", ["page-1-block-07"])},
        ],
        "experience": [
            {
                "employer": make_field("TechCorp", ["page-1-block-09"]),
                "role": make_field("Senior Backend Engineer", ["page-1-block-09"]),
                "location": None,
                "start_date": make_field("2021", ["page-1-block-09"]),
                "end_date": None,
                "current_role": make_field(True, ["page-1-block-09"]),
                "responsibilities": [
                    make_field("Architected high-throughput REST APIs handling 50M daily requests.", ["page-1-block-10"]),
                    make_field("Optimized database queries, reducing API latency by 35%.", ["page-1-block-11"]),
                ],
                "technologies": [make_field("Python", ["page-1-block-10"])],
            },
            {
                "employer": make_field("DataFlow Inc", ["page-1-block-12"]),
                "role": make_field("Software Engineer", ["page-1-block-12"]),
                "location": None,
                "start_date": make_field("2018", ["page-1-block-12"]),
                "end_date": make_field("2021", ["page-1-block-12"]),
                "current_role": make_field(False, ["page-1-block-12"]),
                "responsibilities": [make_field("Built real-time telemetry streaming pipelines using Kafka and Python.", ["page-1-block-13"])],
                "technologies": [make_field("Kafka", ["page-1-block-13"]), make_field("Python", ["page-1-block-13"])],
            },
        ],
        "projects": [],
        "education": [
            {
                "institution": make_field("University of California, Berkeley", ["page-1-block-15"]),
                "degree": make_field("B.S. in Computer Science", ["page-1-block-15"]),
                "field": make_field("Computer Science", ["page-1-block-15"]),
                "start_date": make_field("2014", ["page-1-block-15"]),
                "end_date": make_field("2018", ["page-1-block-15"]),
                "grade": None,
                "location": None,
            }
        ],
        "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    2: {
        "fixture_meta": {"fixture_id": 2, "fixture_slug": "02_two_column", "filename": "02_two_column.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Elena Rostova", ["page-1-block-01"]),
            "email": make_field("elena.rostova@example.com", ["page-1-block-02"]),
            "phone": None,
            "location": make_field("Boston, MA", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": None,
        "target_role": make_field("Lead Data Scientist", ["page-1-block-02"]),
        "skills": [
            {"name": make_field("Python", ["page-1-block-13"]), "category": make_field("Languages", ["page-1-block-13"])},
            {"name": make_field("PyTorch", ["page-1-block-14"]), "category": make_field("ML Frameworks", ["page-1-block-14"])},
        ],
        "experience": [
            {
                "employer": make_field("BioTech AI Labs", ["page-1-block-04"]),
                "role": make_field("Lead Data Scientist", ["page-1-block-04"]),
                "location": None, "start_date": make_field("2020", ["page-1-block-04"]), "end_date": None, "current_role": make_field(True, ["page-1-block-04"]),
                "responsibilities": [make_field("Developed deep learning models for genomic sequence analysis.", ["page-1-block-05"])],
                "technologies": []
            }
        ],
        "projects": [],
        "education": [
            {
                "institution": make_field("MIT", ["page-1-block-10"]),
                "degree": make_field("Ph.D. in Computational Biology", ["page-1-block-10"]),
                "field": make_field("Computational Biology", ["page-1-block-10"]),
                "start_date": make_field("2012", ["page-1-block-10"]), "end_date": make_field("2017", ["page-1-block-10"]), "grade": None, "location": None
            }
        ],
        "certifications": [{"name": make_field("AWS Certified Machine Learning - Specialty", ["page-1-block-17"])}],
        "licenses": [], "achievements": [],
        "publications": [{"title": make_field("Genomic Sequence Transformer Networks", ["page-1-block-19"])}],
        "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    3: {
        "fixture_meta": {"fixture_id": 3, "fixture_slug": "03_sidebar", "filename": "03_sidebar.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Marcus Vance", ["page-1-block-12"]),
            "email": make_field("marcus.vance@example.com", ["page-1-block-02"]),
            "phone": make_field("+1 (555) 019-4820", ["page-1-block-03"]),
            "location": make_field("Austin, TX", ["page-1-block-04"]),
            "linkedin": make_field("linkedin.com/in/mvance", ["page-1-block-05"]),
            "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Principal Cloud Infrastructure Systems Engineer with over 10 years experience designing high-availability distributed systems.", ["page-1-block-16"]),
        "target_role": make_field("Principal Cloud Infrastructure Systems Engineer", ["page-1-block-13"]),
        "skills": [
            {"name": make_field("Python", ["page-1-block-07"]), "category": None},
            {"name": make_field("AWS", ["page-1-block-09"]), "category": None}
        ],
        "experience": [
            {
                "employer": make_field("CloudScale Systems", ["page-1-block-18"]),
                "role": make_field("Principal Cloud Engineer", ["page-1-block-18"]),
                "location": None, "start_date": make_field("2020", ["page-1-block-18"]), "end_date": None, "current_role": make_field(True, ["page-1-block-18"]),
                "responsibilities": [make_field("Designed Kubernetes orchestration platform automating multiregion deployments.", ["page-1-block-19"])],
                "technologies": []
            }
        ],
        "projects": [],
        "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [],
        "languages": [{"language": make_field("English", ["page-1-block-11"])}],
        "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    4: {
        "fixture_meta": {"fixture_id": 4, "fixture_slug": "04_table_based", "filename": "04_table_based.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Samuel Oak", ["page-1-block-01"]),
            "email": make_field("samuel.oak@example.com", ["page-1-block-02"]),
            "phone": make_field("(555) 321-9876", ["page-1-block-02"]),
            "location": None, "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Experienced Full Stack Developer specializing in React, Node.js, and Python web architectures.", ["page-1-block-04"]),
        "target_role": make_field("Full Stack Developer", ["page-1-block-04"]),
        "skills": [
            {"name": make_field("React", ["page-1-block-06"]), "category": None},
            {"name": make_field("Node.js", ["page-1-block-06"]), "category": None}
        ],
        "experience": [
            {
                "employer": make_field("WebMatrix Inc", ["page-1-block-08"]),
                "role": make_field("Full Stack Lead", ["page-1-block-08"]),
                "location": None, "start_date": make_field("2019", ["page-1-block-08"]), "end_date": None, "current_role": make_field(True, ["page-1-block-08"]),
                "responsibilities": [make_field("Led development of SaaS enterprise dashboard.", ["page-1-block-08"])],
                "technologies": []
            }
        ],
        "projects": [],
        "education": [
            {
                "institution": make_field("Oregon State University", ["page-1-block-10"]),
                "degree": make_field("B.S. Software Engineering", ["page-1-block-10"]),
                "field": make_field("Software Engineering", ["page-1-block-10"]),
                "start_date": make_field("2015", ["page-1-block-10"]), "end_date": make_field("2019", ["page-1-block-10"]), "grade": None, "location": None
            }
        ],
        "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    5: {
        "fixture_meta": {"fixture_id": 5, "fixture_slug": "05_long_multipage", "filename": "05_long_multipage.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Dr. Arthur Pendelton", ["page-1-block-01"]),
            "email": make_field("arthur.pendelton@executive.org", ["page-1-block-02"]),
            "phone": make_field("(555) 999-0001", ["page-1-block-02"]),
            "location": make_field("New York, NY", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Global Engineering Executive with 20+ years of leadership scaling engineering organizations.", ["page-1-block-05"]),
        "target_role": make_field("Vice President of Engineering", ["page-1-block-02"]),
        "skills": [],
        "experience": [
            {
                "employer": make_field("Global Tech Enterprises", ["page-1-block-07"]),
                "role": make_field("VP of Engineering", ["page-1-block-07"]),
                "location": None, "start_date": make_field("2020", ["page-1-block-07"]), "end_date": None, "current_role": make_field(True, ["page-1-block-07"]),
                "responsibilities": [make_field("Oversaw $80M annual engineering budget and managed 450+ headcount.", ["page-1-block-08"])],
                "technologies": []
            }
        ],
        "projects": [],
        "education": [
            {
                "institution": make_field("Stanford University", ["page-3-block-03"]),
                "degree": make_field("Ph.D. in Computer Engineering", ["page-3-block-03"]),
                "field": make_field("Computer Engineering", ["page-3-block-03"]),
                "start_date": make_field("2001", ["page-3-block-03"]), "end_date": make_field("2005", ["page-3-block-03"]), "grade": None, "location": None
            }
        ],
        "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    6: {
        "fixture_meta": {"fixture_id": 6, "fixture_slug": "06_minimal_fresher", "filename": "06_minimal_fresher.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Sam Rivera", ["page-1-block-01"]),
            "email": make_field("sam.rivera@gmail.com", ["page-1-block-02"]),
            "phone": make_field("(555) 444-3322", ["page-1-block-02"]),
            "location": make_field("Chicago, IL", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Enthusiastic Computer Science graduate seeking an entry-level Junior Software Developer position.", ["page-1-block-05"]),
        "target_role": make_field("Junior Software Developer", ["page-1-block-05"]),
        "skills": [{"name": make_field("Python", ["page-1-block-10"]), "category": None}],
        "experience": [],
        "projects": [
            {
                "project_name": make_field("Student Portal App", ["page-1-block-12"]),
                "project_type": make_field("Academic Capstone", ["page-1-block-12"]),
                "description": make_field("Created a web portal using Python Flask and SQLite for managing course assignments.", ["page-1-block-13"]),
                "role": None, "technologies": [make_field("Python", ["page-1-block-13"])], "responsibilities": [], "outcomes": [], "dates": None, "links": [], "evidence_block_ids": ["page-1-block-13"], "confidence": "HIGH"
            }
        ],
        "education": [
            {
                "institution": make_field("University of Illinois Chicago", ["page-1-block-07"]),
                "degree": make_field("B.S. in Computer Science", ["page-1-block-07"]),
                "field": make_field("Computer Science", ["page-1-block-07"]),
                "start_date": None, "end_date": make_field("2024-05", ["page-1-block-07"]), "grade": make_field("3.7 / 4.0", ["page-1-block-08"]), "location": None
            }
        ],
        "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    7: {
        "fixture_meta": {"fixture_id": 7, "fixture_slug": "07_senior_technical", "filename": "07_senior_technical.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Vikram Patel", ["page-1-block-01"]),
            "email": make_field("vikram.patel@techleader.io", ["page-1-block-02"]),
            "phone": None,
            "location": make_field("Seattle, WA", ["page-1-block-02"]),
            "linkedin": None, "github": make_field("github.com/vpatel-staff", ["page-1-block-02"]), "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Staff Distributed Systems Architect with 12+ years of experience designing fault-tolerant cloud infrastructures.", ["page-1-block-05"]),
        "target_role": make_field("Staff Distributed Systems Architect", ["page-1-block-02"]),
        "skills": [{"name": make_field("Go", ["page-1-block-07"]), "category": make_field("Languages", ["page-1-block-07"])}],
        "experience": [
            {
                "employer": make_field("HyperScale Systems", ["page-1-block-11"]),
                "role": make_field("Staff Engineer & Architect", ["page-1-block-11"]),
                "location": None, "start_date": make_field("2021", ["page-1-block-11"]), "end_date": None, "current_role": make_field(True, ["page-1-block-11"]),
                "responsibilities": [make_field("Lead architect for real-time payment settlement engine processing 100,000 TPS.", ["page-1-block-12"])],
                "technologies": []
            }
        ],
        "projects": [],
        "education": [
            {
                "institution": make_field("University of Washington", ["page-1-block-18"]),
                "degree": make_field("M.S. in Computer Science", ["page-1-block-18"]),
                "field": make_field("Computer Science", ["page-1-block-18"]),
                "start_date": make_field("2012", ["page-1-block-18"]), "end_date": make_field("2014", ["page-1-block-18"]), "grade": None, "location": None
            }
        ],
        "certifications": [{"name": make_field("Certified Kubernetes Administrator (CKA)", ["page-1-block-19"])}],
        "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    8: {
        "fixture_meta": {"fixture_id": 8, "fixture_slug": "08_career_change", "filename": "08_career_change.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Jordan Lee", ["page-1-block-01"]),
            "email": make_field("jordan.lee@email.com", ["page-1-block-02"]),
            "phone": None, "location": make_field("New York, NY", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Former Financial Analyst transitioning into Software Engineering.", ["page-1-block-05"]),
        "target_role": make_field("Full Stack Developer", ["page-1-block-02"]),
        "skills": [{"name": make_field("Python", ["page-1-block-07"]), "category": None}],
        "experience": [
            {
                "employer": make_field("WallStreet Capital", ["page-1-block-13"]),
                "role": make_field("Senior Financial Analyst", ["page-1-block-13"]),
                "location": None, "start_date": make_field("2018", ["page-1-block-13"]), "end_date": make_field("2023", ["page-1-block-13"]), "current_role": make_field(False, ["page-1-block-13"]),
                "responsibilities": [make_field("Built automated Python financial reporting scripts, cutting monthly close time by 40 hours.", ["page-1-block-14"])],
                "technologies": [make_field("Python", ["page-1-block-14"])]
            }
        ],
        "projects": [
            {
                "project_name": make_field("Portfolio Analytics Dashboard", ["page-1-block-10"]),
                "project_type": None, "description": make_field("Built React frontend with Recharts visualization and Python FastAPI backend.", ["page-1-block-11"]),
                "role": None, "technologies": [make_field("React", ["page-1-block-11"]), make_field("FastAPI", ["page-1-block-11"])], "responsibilities": [], "outcomes": [], "dates": None, "links": [], "evidence_block_ids": ["page-1-block-11"], "confidence": "HIGH"
            }
        ],
        "education": [
            {
                "institution": make_field("NYU Stern School of Business", ["page-1-block-17"]),
                "degree": make_field("B.S. in Finance", ["page-1-block-17"]),
                "field": make_field("Finance", ["page-1-block-17"]),
                "start_date": make_field("2014", ["page-1-block-17"]), "end_date": make_field("2018", ["page-1-block-17"]), "grade": None, "location": None
            }
        ],
        "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    9: {
        "fixture_meta": {"fixture_id": 9, "fixture_slug": "09_academic_cv", "filename": "09_academic_cv.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Dr. Catherine Thorne", ["page-1-block-01"]),
            "email": make_field("catherine.thorne@university.edu", ["page-1-block-02"]),
            "phone": None, "location": make_field("Cambridge, MA", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": None,
        "target_role": make_field("Associate Professor of Computer Science", ["page-1-block-02"]),
        "skills": [], "experience": [], "projects": [],
        "education": [
            {
                "institution": make_field("Harvard University", ["page-1-block-07"]),
                "degree": make_field("Ph.D. in Computer Science", ["page-1-block-07"]),
                "field": make_field("Computer Science", ["page-1-block-07"]),
                "start_date": None, "end_date": make_field("2015", ["page-1-block-07"]), "grade": None, "location": None
            }
        ],
        "certifications": [], "licenses": [], "achievements": [],
        "publications": [{"title": make_field("Neural Program Synthesis with Constraint Verification", ["page-2-block-03"])}],
        "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    10: {
        "fixture_meta": {"fixture_id": 10, "fixture_slug": "10_project_heavy", "filename": "10_project_heavy.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("David Kim", ["page-1-block-01"]),
            "email": make_field("david.kim@dev.io", ["page-1-block-02"]),
            "phone": None, "location": None, "github": make_field("github.com/dkim-dev", ["page-1-block-02"]), "linkedin": None, "portfolio": None, "other_links": []
        },
        "professional_summary": None, "target_role": make_field("Open Source Contributor & Software Developer", ["page-1-block-02"]),
        "skills": [{"name": make_field("Rust", ["page-1-block-05"]), "category": None}],
        "experience": [],
        "projects": [
            {
                "project_name": make_field("FastGraphDB", ["page-1-block-08"]),
                "project_type": make_field("Graph Database Engine", ["page-1-block-08"]),
                "description": make_field("Rust-based graph storage engine supporting Cypher queries.", ["page-1-block-09"]),
                "role": None, "technologies": [make_field("Rust", ["page-1-block-10"])], "responsibilities": [], "outcomes": [], "dates": None, "links": [], "evidence_block_ids": ["page-1-block-09"], "confidence": "HIGH"
            }
        ],
        "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    11: {
        "fixture_meta": {"fixture_id": 11, "fixture_slug": "11_freelance", "filename": "11_freelance.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Maya Lin", ["page-1-block-01"]),
            "email": make_field("maya@linconsulting.com", ["page-1-block-02"]),
            "phone": None, "location": make_field("San Jose, CA", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": None, "target_role": make_field("Independent Cloud Solutions Architect & Consultant", ["page-1-block-02"]),
        "skills": [],
        "experience": [
            {
                "employer": make_field("FinTech Client Alpha", ["page-1-block-07"]),
                "role": make_field("Senior Cloud Consultant", ["page-1-block-07"]),
                "location": None, "start_date": make_field("2023", ["page-1-block-07"]), "end_date": make_field("2023", ["page-1-block-07"]), "current_role": make_field(False, ["page-1-block-07"]),
                "responsibilities": [make_field("Migrated monolithic payment backend to AWS Lambda and DynamoDB.", ["page-1-block-08"])],
                "technologies": []
            }
        ],
        "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    12: {
        "fixture_meta": {"fixture_id": 12, "fixture_slug": "12_multiple_roles", "filename": "12_multiple_roles.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Chris Evans", ["page-1-block-01"]),
            "email": make_field("chris.evans@email.com", ["page-1-block-02"]),
            "phone": None, "location": make_field("Austin, TX", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": None, "target_role": make_field("Staff Software Engineer", ["page-1-block-02"]),
        "skills": [],
        "experience": [
            {
                "employer": make_field("TechCorp Solutions", ["page-1-block-05"]),
                "role": make_field("Staff Software Engineer", ["page-1-block-06"]),
                "location": None, "start_date": make_field("2022", ["page-1-block-06"]), "end_date": None, "current_role": make_field(True, ["page-1-block-06"]),
                "responsibilities": [make_field("Promoted to Staff Engineer to drive technical roadmap.", ["page-1-block-07"])],
                "technologies": []
            },
            {
                "employer": make_field("TechCorp Solutions", ["page-1-block-05"]),
                "role": make_field("Senior Software Engineer", ["page-1-block-09"]),
                "location": None, "start_date": make_field("2020", ["page-1-block-09"]), "end_date": make_field("2022", ["page-1-block-09"]), "current_role": make_field(False, ["page-1-block-09"]),
                "responsibilities": [make_field("Promoted from Engineer II after leading successful migration.", ["page-1-block-10"])],
                "technologies": []
            }
        ],
        "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    13: {
        "fixture_meta": {"fixture_id": 13, "fixture_slug": "13_overlapping_dates", "filename": "13_overlapping_dates.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Taylor Morgan", ["page-1-block-01"]),
            "email": make_field("taylor.morgan@email.com", ["page-1-block-02"]),
            "phone": None, "location": make_field("Denver, CO", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": None, "target_role": make_field("Full Stack Engineer & Tech Advisor", ["page-1-block-02"]),
        "skills": [],
        "experience": [
            {
                "employer": make_field("Apex Innovations", ["page-1-block-05"]),
                "role": make_field("Full Stack Lead Engineer", ["page-1-block-05"]),
                "location": None, "start_date": make_field("2021-01", ["page-1-block-05"]), "end_date": None, "current_role": make_field(True, ["page-1-block-05"]),
                "responsibilities": [make_field("Leading full stack development team building SaaS platform.", ["page-1-block-06"])],
                "technologies": []
            },
            {
                "employer": make_field("Startup Accelerator Project", ["page-1-block-07"]),
                "role": make_field("Technical Co-Founder & Advisor", ["page-1-block-07"]),
                "location": None, "start_date": make_field("2021-03", ["page-1-block-07"]), "end_date": make_field("2022-12", ["page-1-block-07"]), "current_role": make_field(False, ["page-1-block-07"]),
                "responsibilities": [make_field("Concurrent co-founder role building MVP web app.", ["page-1-block-08"])],
                "technologies": []
            }
        ],
        "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    14: {
        "fixture_meta": {"fixture_id": 14, "fixture_slug": "14_unusual_headings", "filename": "14_unusual_headings.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Morgan Reed", ["page-1-block-01"]),
            "email": make_field("morgan.reed@creative.net", ["page-1-block-02"]),
            "phone": None, "location": None, "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": None, "target_role": make_field("Creative Developer & Designer", ["page-1-block-02"]),
        "skills": [{"name": make_field("JavaScript", ["page-1-block-09"]), "category": None}],
        "experience": [
            {
                "employer": make_field("PixelCraft Studios", ["page-1-block-05"]),
                "role": make_field("Lead Interactive Developer", ["page-1-block-05"]),
                "location": None, "start_date": make_field("2020", ["page-1-block-05"]), "end_date": None, "current_role": make_field(True, ["page-1-block-05"]),
                "responsibilities": [make_field("Designed interactive 3D WebGL web experiences.", ["page-1-block-06"])],
                "technologies": []
            }
        ],
        "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    15: {
        "fixture_meta": {"fixture_id": 15, "fixture_slug": "15_no_headings", "filename": "15_no_headings.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Riley Harper", ["page-1-block-01"]),
            "email": make_field("riley.harper@email.com", ["page-1-block-02"]),
            "phone": make_field("(555) 777-8899", ["page-1-block-02"]),
            "location": make_field("San Diego, CA", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Software engineer with 5 years experience specializing in web APIs.", ["page-1-block-03"]),
        "target_role": make_field("Software Developer", ["page-1-block-02"]),
        "skills": [{"name": make_field("Python", ["page-1-block-06"]), "category": None}],
        "experience": [
            {
                "employer": make_field("WebWorks Corp", ["page-1-block-04"]),
                "role": make_field("Software Engineer", ["page-1-block-04"]),
                "location": None, "start_date": make_field("2021", ["page-1-block-04"]), "end_date": None, "current_role": make_field(True, ["page-1-block-04"]),
                "responsibilities": [make_field("Developed Python FastAPI services and React components.", ["page-1-block-04"])],
                "technologies": []
            }
        ],
        "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    16: {
        "fixture_meta": {"fixture_id": 16, "fixture_slug": "16_icons", "filename": "16_icons.pdf", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Avery Bennett", ["page-1-block-01"]),
            "email": make_field("avery.bennett@email.com", ["page-1-block-02"]),
            "phone": make_field("(555) 019-7733", ["page-1-block-02"]),
            "location": make_field("Portland, OR", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": make_field("averybennett.dev", ["page-1-block-02"]), "other_links": []
        },
        "professional_summary": None, "target_role": None, "skills": [],
        "experience": [
            {
                "employer": make_field("EcoTech Systems", ["page-1-block-05"]),
                "role": make_field("Senior Backend Developer", ["page-1-block-05"]),
                "location": None, "start_date": make_field("2020", ["page-1-block-05"]), "end_date": None, "current_role": make_field(True, ["page-1-block-05"]),
                "responsibilities": [make_field("Built REST APIs using Python and Django.", ["page-1-block-06"])],
                "technologies": []
            }
        ],
        "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    17: {
        "fixture_meta": {"fixture_id": 17, "fixture_slug": "17_docx_tables", "filename": "17_docx_tables.docx", "expected_status": "SUCCESS"},
        "contact": {
            "full_name": make_field("Marcus Vance", ["page-1-block-01"]),
            "email": make_field("marcus.vance@example.com", ["page-1-block-02"]),
            "phone": make_field("(555) 019-4820", ["page-1-block-02"]),
            "location": make_field("Austin, TX", ["page-1-block-02"]),
            "linkedin": None, "github": None, "portfolio": None, "other_links": []
        },
        "professional_summary": make_field("Principal Systems Engineer with over 10 years experience designing distributed systems.", ["page-1-block-04"]),
        "target_role": make_field("Principal Systems Engineer", ["page-1-block-04"]),
        "skills": [],
        "experience": [
            {
                "employer": make_field("CloudScale Inc.", ["page-1-block-06"]),
                "role": make_field("Lead Architect", ["page-1-block-06"]),
                "location": None, "start_date": make_field("2021", ["page-1-block-06"]), "end_date": None, "current_role": make_field(True, ["page-1-block-06"]),
                "responsibilities": [make_field("Managed team of 8 backend engineers.", ["page-1-block-06"])],
                "technologies": [make_field("Python", ["page-1-block-06"])]
            }
        ],
        "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [], "warnings": [], "unclassified_blocks": []
    },
    18: {
        "fixture_meta": {"fixture_id": 18, "fixture_slug": "18_scanned", "filename": "18_scanned.pdf", "expected_status": "OCR_REQUIRED"},
        "contact": {"full_name": None, "email": None, "phone": None, "location": None, "linkedin": None, "github": None, "portfolio": None, "other_links": []},
        "professional_summary": None, "target_role": None, "skills": [], "experience": [], "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [],
        "warnings": ["Scanned document detected (no selectable text layer); OCR pipeline triggered"],
        "unclassified_blocks": []
    },
    19: {
        "fixture_meta": {"fixture_id": 19, "fixture_slug": "19_poor_ocr", "filename": "19_poor_ocr.pdf", "expected_status": "OCR_POOR"},
        "contact": {"full_name": None, "email": None, "phone": None, "location": None, "linkedin": None, "github": None, "portfolio": None, "other_links": []},
        "professional_summary": None, "target_role": None, "skills": [], "experience": [], "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [],
        "warnings": ["Low confidence OCR text extraction due to image noise/blur"],
        "unclassified_blocks": []
    },
    20: {
        "fixture_meta": {"fixture_id": 20, "fixture_slug": "20_empty", "filename": "20_empty.pdf", "expected_status": "EMPTY_FILE"},
        "contact": {"full_name": None, "email": None, "phone": None, "location": None, "linkedin": None, "github": None, "portfolio": None, "other_links": []},
        "professional_summary": None, "target_role": None, "skills": [], "experience": [], "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [],
        "warnings": ["File content is zero bytes"],
        "unclassified_blocks": []
    },
    21: {
        "fixture_meta": {"fixture_id": 21, "fixture_slug": "21_corrupted", "filename": "21_corrupted.pdf", "expected_status": "CORRUPTED_FILE"},
        "contact": {"full_name": None, "email": None, "phone": None, "location": None, "linkedin": None, "github": None, "portfolio": None, "other_links": []},
        "professional_summary": None, "target_role": None, "skills": [], "experience": [], "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [],
        "warnings": ["File structure is malformed or unparseable"],
        "unclassified_blocks": []
    },
    22: {
        "fixture_meta": {"fixture_id": 22, "fixture_slug": "22_encrypted", "filename": "22_encrypted.pdf", "expected_status": "ENCRYPTED_FILE"},
        "contact": {"full_name": None, "email": None, "phone": None, "location": None, "linkedin": None, "github": None, "portfolio": None, "other_links": []},
        "professional_summary": None, "target_role": None, "skills": [], "experience": [], "projects": [], "education": [], "certifications": [], "licenses": [], "achievements": [], "publications": [], "languages": [], "volunteer_experience": [], "training": [], "links": [], "additional_sections": [],
        "warnings": ["File is password encrypted"],
        "unclassified_blocks": []
    },
}

# ==============================================================================
# Fixture Specs Registry
# ==============================================================================

FIXTURES: List[FixtureSpec] = [
    FixtureSpec(1, "01_single_column", "01_single_column.pdf", "01_single_column.json", "pdf", "Single column PDF", gen_01_single_column, GOLDEN_BENCHMARKS[1]),
    FixtureSpec(2, "02_two_column", "02_two_column.pdf", "02_two_column.json", "pdf", "Two column PDF", gen_02_two_column, GOLDEN_BENCHMARKS[2]),
    FixtureSpec(3, "03_sidebar", "03_sidebar.pdf", "03_sidebar.json", "pdf", "Sidebar layout PDF", gen_03_sidebar, GOLDEN_BENCHMARKS[3]),
    FixtureSpec(4, "04_table_based", "04_table_based.pdf", "04_table_based.json", "pdf", "Table based PDF", gen_04_table_based, GOLDEN_BENCHMARKS[4]),
    FixtureSpec(5, "05_long_multipage", "05_long_multipage.pdf", "05_long_multipage.json", "pdf", "Long multipage PDF", gen_05_long_multipage, GOLDEN_BENCHMARKS[5]),
    FixtureSpec(6, "06_minimal_fresher", "06_minimal_fresher.pdf", "06_minimal_fresher.json", "pdf", "Minimal fresher PDF", gen_06_minimal_fresher, GOLDEN_BENCHMARKS[6]),
    FixtureSpec(7, "07_senior_technical", "07_senior_technical.pdf", "07_senior_technical.json", "pdf", "Senior technical PDF", gen_07_senior_technical, GOLDEN_BENCHMARKS[7]),
    FixtureSpec(8, "08_career_change", "08_career_change.pdf", "08_career_change.json", "pdf", "Career change PDF", gen_08_career_change, GOLDEN_BENCHMARKS[8]),
    FixtureSpec(9, "09_academic_cv", "09_academic_cv.pdf", "09_academic_cv.json", "pdf", "Academic CV PDF", gen_09_academic_cv, GOLDEN_BENCHMARKS[9]),
    FixtureSpec(10, "10_project_heavy", "10_project_heavy.pdf", "10_project_heavy.json", "pdf", "Project heavy PDF", gen_10_project_heavy, GOLDEN_BENCHMARKS[10]),
    FixtureSpec(11, "11_freelance", "11_freelance.pdf", "11_freelance.json", "pdf", "Freelance consultant PDF", gen_11_freelance, GOLDEN_BENCHMARKS[11]),
    FixtureSpec(12, "12_multiple_roles", "12_multiple_roles.pdf", "12_multiple_roles.json", "pdf", "Multiple roles same employer PDF", gen_12_multiple_roles, GOLDEN_BENCHMARKS[12]),
    FixtureSpec(13, "13_overlapping_dates", "13_overlapping_dates.pdf", "13_overlapping_dates.json", "pdf", "Overlapping dates PDF", gen_13_overlapping_dates, GOLDEN_BENCHMARKS[13]),
    FixtureSpec(14, "14_unusual_headings", "14_unusual_headings.pdf", "14_unusual_headings.json", "pdf", "Unusual section headings PDF", gen_14_unusual_headings, GOLDEN_BENCHMARKS[14]),
    FixtureSpec(15, "15_no_headings", "15_no_headings.pdf", "15_no_headings.json", "pdf", "No section headings PDF", gen_15_no_headings, GOLDEN_BENCHMARKS[15]),
    FixtureSpec(16, "16_icons", "16_icons.pdf", "16_icons.json", "pdf", "Unicode icons PDF", gen_16_icons, GOLDEN_BENCHMARKS[16]),
    FixtureSpec(17, "17_docx_tables", "17_docx_tables.docx", "17_docx_tables.json", "docx", "DOCX tables resume", gen_17_docx_tables, GOLDEN_BENCHMARKS[17]),
    FixtureSpec(18, "18_scanned", "18_scanned.pdf", "18_scanned.json", "pdf", "Scanned raster image PDF", gen_18_scanned, GOLDEN_BENCHMARKS[18]),
    FixtureSpec(19, "19_poor_ocr", "19_poor_ocr.pdf", "19_poor_ocr.json", "pdf", "Poor quality OCR PDF", gen_19_poor_ocr, GOLDEN_BENCHMARKS[19]),
    FixtureSpec(20, "20_empty", "20_empty.pdf", "20_empty.json", "pdf", "Zero byte empty PDF", gen_20_empty, GOLDEN_BENCHMARKS[20]),
    FixtureSpec(21, "21_corrupted", "21_corrupted.pdf", "21_corrupted.json", "pdf", "Corrupted PDF header", gen_21_corrupted, GOLDEN_BENCHMARKS[21]),
    FixtureSpec(22, "22_encrypted", "22_encrypted.pdf", "22_encrypted.json", "pdf", "Encrypted PDF", gen_22_encrypted, GOLDEN_BENCHMARKS[22]),
]


def run_generator(force: bool = False, quiet: bool = False, golden_only: bool = False, files_only: bool = False, selected_ids: Optional[List[int]] = None) -> int:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    GOLDEN_DIR.mkdir(parents=True, exist_ok=True)

    targets = FIXTURES
    if selected_ids:
        targets = [f for f in FIXTURES if f.id in selected_ids]

    success_count = 0
    skip_count = 0
    fail_count = 0

    for spec in targets:
        fixture_exists = spec.fixture_path.exists()
        golden_exists = spec.golden_path.exists()

        if not force and fixture_exists and golden_exists:
            if not quiet:
                logger.info(f"[SKIP] {spec.slug}: Output files already exist. Use --force to overwrite.")
            skip_count += 1
            continue

        try:
            if not golden_only:
                if not quiet:
                    logger.info(f"Generating fixture document: {spec.filename}...")
                spec.generator_fn(spec.fixture_path, spec.golden_data)

            if not files_only:
                if not quiet:
                    logger.info(f"Generating golden JSON: {spec.golden_filename}...")
                normalized_data = normalize_golden_benchmark(spec.golden_data)
                ParsedResumeSchema.model_validate(normalized_data)
                with open(spec.golden_path, "w", encoding="utf-8") as f:
                    json.dump(normalized_data, f, indent=2)

            if not quiet:
                logger.info(f"[OK] {spec.slug}: Successfully generated.")
            success_count += 1
        except Exception as e:
            logger.error(f"[ERROR] {spec.slug}: Generation failed — {e}", exc_info=True)
            fail_count += 1

    if not quiet:
        print("\n" + "=" * 70)
        print("Synthetic Resume Fixture Generation Summary")
        print("=" * 70)
        print(f"Total Fixtures Target : {len(targets)}")
        print(f"Successfully Created  : {success_count}")
        print(f"Skipped (Existing)    : {skip_count}")
        print(f"Failed                : {fail_count}")
        print(f"Fixture Directory     : {FIXTURES_DIR}")
        print(f"Golden Directory      : {GOLDEN_DIR}")
        print("=" * 70)

    return 0 if fail_count == 0 else 1


def clean_generated_files() -> int:
    removed = 0
    for spec in FIXTURES:
        if spec.fixture_path.exists():
            spec.fixture_path.unlink()
            removed += 1
        if spec.golden_path.exists():
            spec.golden_path.unlink()
            removed += 1
    logger.info(f"Cleaned {removed} generated fixture and golden files.")
    return 0


def main() -> None:
    parser = argparse.ArgumentParser(description="Synthetic Resume Fixture Generator CLI")
    parser.add_argument("-f", "--force", action="store_true", help="Force overwrite of existing fixture documents and golden JSON files.")
    parser.add_argument("-i", "--fixture", action="append", help="Generate specific fixture by ID (1-22) or slug.")
    parser.add_argument("-l", "--list", action="store_true", help="List all registered fixture specifications.")
    parser.add_argument("--clean", action="store_true", help="Remove generated fixture files and exit.")
    parser.add_argument("--golden-only", action="store_true", help="Generate golden JSON files only.")
    parser.add_argument("--files-only", action="store_true", help="Generate resume document files only.")
    parser.add_argument("-q", "--quiet", action="store_true", help="Suppress detailed per-fixture logging.")

    args = parser.parse_args()

    if args.list:
        print("\nRegistered Fixture Specifications (1..22):")
        print("-" * 70)
        for spec in FIXTURES:
            print(f"[{spec.id:02d}] {spec.slug:<22} | Format: {spec.file_type:<4} | {spec.description}")
        sys.exit(0)

    if args.clean:
        sys.exit(clean_generated_files())

    selected_ids = None
    if args.fixture:
        selected_ids = []
        for item in args.fixture:
            if item.isdigit():
                selected_ids.append(int(item))
            else:
                for f in FIXTURES:
                    if f.slug == item or f.filename == item:
                        selected_ids.append(f.id)

    exit_code = run_generator(
        force=args.force,
        quiet=args.quiet,
        golden_only=args.golden_only,
        files_only=args.files_only,
        selected_ids=selected_ids,
    )
    sys.exit(exit_code)


if __name__ == "__main__":
    main()
