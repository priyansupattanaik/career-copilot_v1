import asyncio
import io
import json
import uuid

import httpx
from docx import Document
from supabase import ClientOptions, create_client

from app.config import get_settings
from app.main import app
from app.supabase_clients import create_admin_supabase_client

TABLES = [
    "profiles",
    "candidate_preferences",
    "candidate_skills",
    "candidate_experiences",
    "candidate_projects",
    "candidate_education",
    "candidate_certifications",
    "candidate_languages",
    "candidate_links",
    "resumes",
    "resume_versions",
    "job_descriptions",
    "ats_analyses",
    "ats_evidence",
    "resume_suggestions",
    "resume_exports",
    "interview_sessions",
    "interview_questions",
    "interview_responses",
    "interview_reports",
    "learning_paths",
    "learning_items",
    "learning_resources",
    "jobs",
    "job_recommendations",
    "saved_jobs",
    "notification_preferences",
    "privacy_preferences",
    "activity_events",
    "user_notifications",
]
BUCKETS = {"candidate-documents", "candidate-avatars", "interview-media"}


def user_client(settings, token):
    client = create_client(
        settings.supabase_url,
        settings.supabase_publishable_key,
        options=ClientOptions(headers={"Authorization": f"Bearer {token}"}),
    )
    client.postgrest.auth(token)
    return client


def make_docx():
    document = Document()
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python, SQL")
    document.add_heading("Experience", level=1)
    document.add_paragraph("Built evidence-led product analytics workflows.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


async def main():
    settings = get_settings()
    admin = create_admin_supabase_client(settings)
    suffix = uuid.uuid4().hex
    credentials = [
        (f"codex-rls-a-{suffix}@example.invalid", uuid.uuid4().hex + "A1!"),
        (f"codex-rls-b-{suffix}@example.invalid", uuid.uuid4().hex + "B1!"),
    ]
    user_ids = []
    storage_paths = []
    summary = {}
    try:
        for table in TABLES:
            admin.table(table).select("*", count="exact", head=True).execute()
        summary["tables"] = f"{len(TABLES)}/{len(TABLES)}"

        buckets = {bucket.name for bucket in admin.storage.list_buckets()}
        assert BUCKETS <= buckets
        summary["buckets"] = f"{len(BUCKETS)}/{len(BUCKETS)}"

        clients = []
        tokens = []
        for email, password in credentials:
            created = admin.auth.admin.create_user(
                {"email": email, "password": password, "email_confirm": True}
            )
            user_ids.append(str(created.user.id))
            browser = create_client(settings.supabase_url, settings.supabase_publishable_key)
            session = browser.auth.sign_in_with_password({"email": email, "password": password})
            tokens.append(session.session.access_token)
            clients.append(user_client(settings, session.session.access_token))

        a, b = clients
        a_id, b_id = user_ids
        a.table("profiles").update({"full_name": "RLS User A", "current_role": "Analyst"}).eq(
            "id", a_id
        ).execute()
        b.table("profiles").update({"full_name": "RLS User B"}).eq("id", b_id).execute()
        assert {row["id"] for row in a.table("profiles").select("id").execute().data} == {a_id}
        assert a.table("profiles").select("id").eq("id", b_id).execute().data == []
        assert a.table("profiles").update({"full_name": "blocked"}).eq("id", b_id).execute().data == []
        skill = (
            a.table("candidate_skills")
            .insert({"user_id": a_id, "name": "SQL", "normalized_name": "sql"})
            .execute()
            .data[0]
        )
        assert b.table("candidate_skills").select("id").eq("id", skill["id"]).execute().data == []
        summary["database_rls"] = "passed"

        storage_path = f"{a_id}/rls-check/{uuid.uuid4()}.pdf"
        storage_paths.append(storage_path)
        a.storage.from_(settings.document_bucket).upload(
            storage_path, b"%PDF-1.4\n%%EOF", {"content-type": "application/pdf", "upsert": "false"}
        )
        assert a.storage.from_(settings.document_bucket).download(storage_path)
        blocked = False
        try:
            b.storage.from_(settings.document_bucket).download(storage_path)
        except Exception:
            blocked = True
        assert blocked
        summary["storage_rls"] = "passed"

        transport = httpx.ASGITransport(app=app, raise_app_exceptions=False)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as api:
            headers = {"Authorization": f"Bearer {tokens[0]}"}
            profile = await api.get("/api/v1/profile", headers=headers)
            assert profile.status_code == 200, profile.text
            jd = await api.post(
                "/api/v1/job-descriptions",
                headers=headers,
                json={
                    "title": "Live integration role",
                    "company": "Integration Test",
                    "role_title": "Analyst",
                    "raw_text": "Requirements\nPython\nSQL\nEvidence-led product analysis and communication.",
                },
            )
            assert jd.status_code == 201, jd.text
            resume = await api.post(
                "/api/v1/resumes",
                headers=headers,
                data={"title": "Live integration resume"},
                files={
                    "file": (
                        "integration-resume.docx",
                        make_docx(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            assert resume.status_code == 201, resume.text
            resume_body = resume.json()
            storage_paths.append(resume_body["version"]["storage_path"])
            confirm = await api.post(
                f"/api/v1/resume-versions/{resume_body['version']['id']}/confirm", headers=headers
            )
            assert confirm.status_code == 200, confirm.text
            summary["backend_live_flow"] = "passed"

        print(json.dumps(summary, sort_keys=True))
    finally:
        if storage_paths:
            try:
                admin.storage.from_(settings.document_bucket).remove(storage_paths)
            except Exception:
                pass
        for user_id in user_ids:
            try:
                admin.auth.admin.delete_user(user_id)
            except Exception:
                pass
        print("temporary_cleanup=complete")


if __name__ == "__main__":
    asyncio.run(main())
