"""
Plain-text extraction from PDF/DOCX.

PDF: Docling only (layout-aware). No pypdf fallback — install docling when missing.
DOCX: python-docx (native structure) with optional Docling when useful.
"""

from __future__ import annotations

import io
import logging
import re
import tempfile
from pathlib import Path

from docx import Document

from app.core.errors import ApiError

logger = logging.getLogger(__name__)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _normalize_extracted_text(text: str) -> str:
    """Preserve line structure; collapse weird whitespace without merging sections."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00ad", "").replace("\ufeff", "")
    lines: list[str] = []
    for raw in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw).strip()
        # Drop markdown heading markers Docling may emit — keep the heading words.
        if line.startswith("#"):
            line = re.sub(r"^#+\s*", "", line).strip()
        lines.append(line)
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if not line:
            blank_run += 1
            if blank_run <= 1:
                cleaned.append("")
            continue
        blank_run = 0
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def _require_docling():
    try:
        from docling.document_converter import DocumentConverter  # type: ignore

        return DocumentConverter
    except Exception as exc:
        raise ApiError(
            503,
            "docling_not_installed",
            "Docling is required for accurate resume PDF parsing. "
            "Install it with: pip install 'docling>=2.0,<3'",
        ) from exc


def _extract_with_docling(content: bytes, suffix: str) -> str:
    DocumentConverter = _require_docling()
    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
            handle.write(content)
            tmp_path = handle.name
        converter = DocumentConverter()
        result = converter.convert(tmp_path)
        doc = result.document
        text = ""
        if hasattr(doc, "export_to_markdown"):
            text = doc.export_to_markdown() or ""
        elif hasattr(doc, "export_to_text"):
            text = doc.export_to_text() or ""
        else:
            text = str(doc)
        text = _normalize_extracted_text(text)
        if not text:
            raise ApiError(
                422,
                "document_has_no_text",
                "Docling found no usable text. Scanned image PDFs need OCR-enabled Docling models.",
            )
        return text
    except ApiError:
        raise
    except Exception as exc:
        logger.exception("docling_extract_failed")
        raise ApiError(
            400,
            "document_parse_failed",
            "Docling could not parse this document. Try a text-based PDF or DOCX.",
        ) from exc
    finally:
        if tmp_path:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                pass


def _docx_paragraph_text(document: Document) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            for cell in cells:
                for part in cell.splitlines():
                    part = part.strip()
                    if part:
                        lines.append(part)
    return lines


def extract_text(content: bytes, mime_type: str) -> str:
    """
    Extract plain text from PDF or DOCX bytes.

    PDF uses Docling only (accurate layout). No silent pypdf fallback.
    """
    if not content:
        raise ApiError(400, "empty_document", "The selected document is empty.")

    if mime_type == PDF_MIME:
        return _extract_with_docling(content, ".pdf")

    if mime_type == DOCX_MIME:
        # Prefer Docling for consistency when available; native docx is reliable for Word files.
        try:
            return _extract_with_docling(content, ".docx")
        except ApiError as exc:
            if exc.code == "docling_not_installed":
                document = Document(io.BytesIO(content))
                text = _normalize_extracted_text("\n".join(_docx_paragraph_text(document)))
                if not text:
                    raise ApiError(422, "document_has_no_text", "No usable text was found in the DOCX.")
                return text
            # If Docling is installed but failed on this file, try native docx once.
            if exc.code == "document_parse_failed":
                document = Document(io.BytesIO(content))
                text = _normalize_extracted_text("\n".join(_docx_paragraph_text(document)))
                if text:
                    return text
            raise

    raise ApiError(415, "unsupported_document_type", "Only PDF and DOCX documents are supported.")
