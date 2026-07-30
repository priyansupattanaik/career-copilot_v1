import hashlib
import io
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from app.errors import ApiError

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ALLOWED_SUFFIXES = {".pdf": PDF_MIME, ".docx": DOCX_MIME}

# Canonical section keys -> accepted heading aliases (normalized lowercase).
HEADING_ALIASES: dict[str, frozenset[str]] = {
    "contact": frozenset(
        {
            "contact",
            "contact details",
            "contact information",
            "personal information",
            "personal details",
            "personal info",
        }
    ),
    "summary": frozenset(
        {
            "summary",
            "profile",
            "objective",
            "professional summary",
            "career summary",
            "executive summary",
            "professional profile",
            "career objective",
            "about me",
            "about",
            "overview",
        }
    ),
    "skills": frozenset(
        {
            "skills",
            "technical skills",
            "core skills",
            "key skills",
            "skill set",
            "skillset",
            "competencies",
            "core competencies",
            "technical competencies",
            "technologies",
            "tech stack",
            "tools",
            "tools technologies",
            "technical proficiencies",
            "areas of expertise",
            "expertise",
        }
    ),
    "experience": frozenset(
        {
            "experience",
            "work experience",
            "professional experience",
            "employment",
            "employment history",
            "work history",
            "career history",
            "professional background",
            "relevant experience",
            "work",
            "career",
            "professional history",
            "internship experience",
            "internships",
        }
    ),
    "projects": frozenset(
        {
            "projects",
            "project",
            "project experience",
            "personal projects",
            "academic projects",
            "key projects",
            "selected projects",
            "project work",
            "notable projects",
            "side projects",
            "portfolio",
            "portfolio projects",
            "major projects",
            "relevant projects",
        }
    ),
    "education": frozenset(
        {
            "education",
            "academic background",
            "academic qualifications",
            "qualifications",
            "academics",
            "educational background",
            "education qualifications",
            "educational qualifications",
            "academic history",
            "degrees",
        }
    ),
    "certifications": frozenset(
        {
            "certifications",
            "certificates",
            "licenses",
            "licenses and certifications",
            "professional certifications",
            "certification",
            "courses",
            "training",
            "trainings",
            "courses and certifications",
        }
    ),
    "languages": frozenset({"languages", "language", "language proficiency", "spoken languages"}),
    "links": frozenset({"links", "profiles", "online profiles", "social links", "websites"}),
    "achievements": frozenset(
        {
            "achievements",
            "accomplishments",
            "awards",
            "honors",
            "awards and achievements",
            "honors and awards",
            "key achievements",
        }
    ),
    "responsibilities": frozenset(
        {
            "responsibilities",
            "key responsibilities",
            "duties",
            "what you will do",
            "role responsibilities",
            "job responsibilities",
        }
    ),
    "requirements": frozenset(
        {
            "requirements",
            "required qualifications",
            "must have",
            "must haves",
            "minimum qualifications",
            "required skills",
            "basic qualifications",
            "essential requirements",
            "what we look for",
            "job requirements",
        }
    ),
    "preferred_qualifications": frozenset(
        {
            "preferred qualifications",
            "preferred skills",
            "nice to have",
            "nice to haves",
            "good to have",
            "bonus skills",
            "desired skills",
            "additional qualifications",
        }
    ),
}

# Longest aliases first so "professional experience" wins over bare fragments.
_HEADING_LOOKUP: list[tuple[str, str]] = sorted(
    ((alias, key) for key, aliases in HEADING_ALIASES.items() for alias in aliases),
    key=lambda item: (-len(item[0]), item[0]),
)

_BULLET_RE = re.compile(r"^[\u2022\u2023\u25E6\u2043\u2219•●○▪▸►\*◦‣⁃\-–—]\s+")
_MONTH = (
    r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
)
_DATE_RANGE_RE = re.compile(
    rf"(?:"
    rf"{_MONTH}\.?\s*'?\d{{2,4}}"
    rf"|(?:19|20)\d{{2}}"
    rf"|\d{{1,2}}[/\-.]\d{{2,4}}"
    rf")"
    rf"\s*[-–—to]+\s*"
    rf"(?:"
    rf"{_MONTH}\.?\s*'?\d{{2,4}}"
    rf"|(?:19|20)\d{{2}}"
    rf"|\d{{1,2}}[/\-.]\d{{2,4}}"
    rf"|present|current|now|ongoing|till\s+date|to\s+date"
    rf")",
    re.I,
)
_EMAIL_RE = re.compile(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?(?:\(?\d{2,4}\)?[\s\-.]?)?\d{3,4}[\s\-.]?\d{3,4}"
)
_URL_RE = re.compile(r"https?://\S+|www\.\S+|(?:linkedin|github)\.com/\S+", re.I)
_ENTRY_SECTIONS = frozenset({"experience", "projects", "education", "certifications", "achievements"})


def safe_filename(filename: str) -> str:
    name = Path(filename).name
    return re.sub(r"[^A-Za-z0-9._-]", "_", name)[:180] or "document"


def sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def validate_document(filename: str, declared_mime: str | None, content: bytes, max_bytes: int) -> str:
    if not content:
        raise ApiError(400, "empty_document", "The selected document is empty.")
    if len(content) > max_bytes:
        raise ApiError(413, "document_too_large", "The selected document exceeds the 10 MB limit.")
    suffix = Path(filename).suffix.lower()
    expected = ALLOWED_SUFFIXES.get(suffix)
    if not expected:
        raise ApiError(415, "unsupported_document_type", "Only PDF and DOCX documents are supported.")
    if declared_mime and declared_mime not in {expected, "application/octet-stream"}:
        raise ApiError(415, "document_mime_mismatch", "The file extension and MIME type do not match.")
    if suffix == ".pdf" and not content.startswith(b"%PDF-"):
        raise ApiError(415, "invalid_pdf_signature", "The selected file is not a valid PDF.")
    if suffix == ".docx":
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                if "word/document.xml" not in archive.namelist():
                    raise ApiError(
                        415, "invalid_docx_structure", "The selected file is not a valid DOCX document."
                    )
        except zipfile.BadZipFile as exc:
            raise ApiError(415, "invalid_docx_archive", "The selected DOCX file is corrupted.") from exc
    return expected


def _docx_paragraph_text(document: Document) -> list[str]:
    """Collect DOCX paragraph text in document order, including simple table cells."""
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            # Keep cell content as separate lines so section headings inside tables still match.
            for cell in cells:
                for part in cell.splitlines():
                    part = part.strip()
                    if part:
                        lines.append(part)
    return lines


def extract_text(content: bytes, mime_type: str) -> str:
    try:
        if mime_type == PDF_MIME:
            reader = PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                raise ApiError(400, "encrypted_pdf", "Password-protected PDFs are not supported.")
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            document = Document(io.BytesIO(content))
            text = "\n".join(_docx_paragraph_text(document))
    except ApiError:
        raise
    except Exception as exc:
        raise ApiError(400, "document_parse_failed", "The document could not be read.") from exc
    if not text.strip():
        raise ApiError(
            422,
            "document_has_no_text",
            "No usable text was found. Scanned documents require OCR, which is not enabled.",
        )
    return text.strip()


def _normalize_heading_label(line: str) -> str:
    cleaned = line.strip().rstrip(":").strip()
    cleaned = re.sub(r"[^a-z0-9\s&/+]", " ", cleaned.lower())
    return re.sub(r"\s+", " ", cleaned).strip()


def _looks_like_heading_line(line: str) -> bool:
    """Headings are short, non-bullet, non-sentence lines (often ALL CAPS or Title Case)."""
    stripped = line.strip()
    if not stripped or len(stripped) > 72:
        return False
    if _BULLET_RE.match(stripped):
        return False
    if _EMAIL_RE.search(stripped) or _URL_RE.search(stripped):
        return False
    # Full sentences are not headings.
    if stripped.endswith(".") and len(stripped.split()) > 4:
        return False
    # Avoid treating long skill dumps as headings.
    if stripped.count(",") >= 3:
        return False
    words = stripped.split()
    if len(words) > 8:
        return False
    return True


def match_section_heading(line: str) -> str | None:
    """Return canonical section key if this line is a known resume/JD heading."""
    if not _looks_like_heading_line(line):
        return None
    normalized = _normalize_heading_label(line)
    if not normalized:
        return None
    for alias, key in _HEADING_LOOKUP:
        if normalized == alias:
            return key
    # Soft match: heading is "SECTION NAME (optional note)" after stripping parentheticals.
    without_parens = re.sub(r"\([^)]*\)", "", normalized).strip()
    without_parens = re.sub(r"\s+", " ", without_parens)
    if without_parens and without_parens != normalized:
        for alias, key in _HEADING_LOOKUP:
            if without_parens == alias:
                return key
    return None


def _is_contact_line(line: str) -> bool:
    if _EMAIL_RE.search(line) or _URL_RE.search(line):
        return True
    digits = sum(ch.isdigit() for ch in line)
    if digits >= 8 and _PHONE_RE.search(line):
        return True
    return False


def _is_bullet_line(line: str) -> bool:
    return bool(_BULLET_RE.match(line.strip()))


def _is_entry_start(line: str, section: str) -> bool:
    """Detect the header line of a new job / project / education entry."""
    if section not in _ENTRY_SECTIONS:
        return False
    stripped = line.strip()
    if not stripped or _is_bullet_line(stripped):
        return False
    if _DATE_RANGE_RE.search(stripped):
        return True
    # Role | Company | Location style headers
    if stripped.count("|") >= 1 and len(stripped) <= 140 and not stripped.endswith("."):
        return True
    # "Role at Company" / "Role - Company"
    if re.search(r"\s+(?:at|@)\s+", stripped, re.I) and len(stripped) <= 120:
        return True
    if re.search(r"\s[-–—]\s", stripped) and len(stripped) <= 120 and not stripped.endswith("."):
        # Prefer title-like lines without trailing sentence punctuation.
        if len(stripped.split()) <= 14:
            return True
    return False


def _group_section_entries(section: str, lines: list[str]) -> list[str]:
    """
    Group experience/projects/education into discrete multi-line entries.
    Each entry is one string with internal newlines so the UI can format cleanly.
    Skills and summary stay as individual lines.
    """
    if not lines:
        return []
    if section not in _ENTRY_SECTIONS:
        # Skills: keep lines; split pure comma lists into cleaner items when short.
        if section == "skills":
            return _normalize_skill_lines(lines)
        return lines

    entries: list[str] = []
    current: list[str] = []

    def flush() -> None:
        nonlocal current
        if current:
            entries.append("\n".join(current).strip())
            current = []

    for line in lines:
        if not line.strip():
            flush()
            continue
        stripped = line.strip()
        # Start a new entry when this line looks like a role/project header and
        # we already have body content (bullets) or a prior header in `current`.
        if current and _is_entry_start(stripped, section) and not _is_bullet_line(stripped):
            prior_has_body = any(_is_bullet_line(item) for item in current)
            prior_is_header = _is_entry_start(current[0], section) or bool(
                _DATE_RANGE_RE.search(current[0])
            )
            # Avoid splitting a single header that spans two short lines.
            if prior_has_body or (prior_is_header and len(current) >= 1):
                # If current is only a header with no bullets yet, still split when
                # the new line is clearly another dated/piped role header.
                if prior_has_body or _DATE_RANGE_RE.search(stripped) or "|" in stripped:
                    flush()
        current.append(stripped)
    flush()
    return entries


def _normalize_skill_lines(lines: list[str]) -> list[str]:
    """Preserve labeled skill rows; expand pure comma lists into readable lines."""
    result: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Keep "Languages: Python, SQL" as one line for readability.
        if ":" in stripped and len(stripped) <= 200:
            result.append(stripped)
            continue
        if stripped.count(",") >= 2 and len(stripped) <= 240 and not _is_bullet_line(stripped):
            parts = [part.strip() for part in re.split(r"[,;|/]", stripped) if part.strip()]
            if 2 <= len(parts) <= 30:
                result.extend(parts)
                continue
        result.append(stripped)
    return result


def extract_sections(text: str, schema_version: str = "resume-extraction-v1") -> dict[str, Any]:
    """
    Parse resume/JD plain text into canonical sections with accurate boundaries.

    Improvements over naive heading matching:
    - Broad heading aliases (Professional Experience, Technical Skills, …)
    - Rejects long sentences that merely contain heading words
    - Auto-collects leading contact lines (email/phone/url)
    - Groups experience/projects/education into separated multi-line entries
    """
    raw_sections: dict[str, list[str]] = {}
    unclassified: list[str] = []
    current: str | None = None
    pending_blank = False

    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            # Preserve blank-line boundaries inside entry sections for grouping.
            if current and current in _ENTRY_SECTIONS and raw_sections.get(current):
                pending_blank = True
            continue

        heading = match_section_heading(line)
        if heading:
            current = heading
            raw_sections.setdefault(current, [])
            pending_blank = False
            continue

        if current is None and _is_contact_line(line):
            raw_sections.setdefault("contact", []).append(line)
            continue

        # Leading name / headline before the first heading → contact block.
        if current is None:
            if not unclassified and len(line) <= 80 and not line.endswith("."):
                unclassified.append(line)
            elif _is_contact_line(line):
                raw_sections.setdefault("contact", []).append(line)
            else:
                unclassified.append(line)
            continue

        if pending_blank:
            # Insert an empty marker so entry grouping can split on blank lines.
            if raw_sections[current] and raw_sections[current][-1] != "":
                raw_sections[current].append("")
            pending_blank = False

        raw_sections[current].append(line)

    # Promote a leading name line into contact when we have other contact data.
    if unclassified and "contact" in raw_sections:
        name_candidate = unclassified[0]
        if len(name_candidate) <= 60 and not _is_contact_line(name_candidate):
            raw_sections["contact"].insert(0, name_candidate)
            unclassified = unclassified[1:]

    sections: dict[str, list[str]] = {}
    for key, lines in raw_sections.items():
        cleaned = [line for line in lines if line is not None]
        grouped = _group_section_entries(key, cleaned)
        if grouped:
            sections[key] = grouped

    warnings: list[str] = []
    if not sections:
        warnings.append("No recognised section headings were found; review all extracted text.")
    else:
        if "experience" not in sections and "projects" not in sections:
            warnings.append("No professional experience or projects section was detected.")
        # Detect possible bleed: projects content inside experience without a projects heading.
        experience_lines = sections.get("experience") or []
        if experience_lines and "projects" not in sections:
            joined = "\n".join(experience_lines).lower()
            if re.search(r"\b(personal projects?|academic projects?|side projects?)\b", joined):
                warnings.append(
                    "Project-like content may still sit under experience; review separation carefully."
                )

    return {
        "schema_version": schema_version,
        "sections": sections,
        "unclassified_blocks": unclassified,
        "warnings": warnings,
        "corrections": {},
    }


def infer_resume_title(filename: str | None) -> str:
    """Derive a resume library title from the uploaded filename."""
    stem = Path(filename or "Resume").stem
    cleaned = re.sub(r"[_\-]+", " ", stem).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return (cleaned or "Resume")[:200]


_ROLE_HINT = re.compile(
    r"\b(engineer|developer|analyst|manager|designer|scientist|architect|specialist|"
    r"lead|intern|consultant|administrator|officer|coordinator|executive|director)\b",
    re.I,
)


def infer_job_metadata(text: str) -> dict[str, str | None]:
    """
    Infer title, role, and company from job-description text so candidates
    do not need to type those fields manually.
    """
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    role: str | None = None
    company: str | None = None
    confidence = "low"

    for line in lines[:60]:
        for label in (
            "job title",
            "position title",
            "role title",
            "designation",
            "opening for",
            "hiring for",
            "title",
            "position",
            "role",
        ):
            match = re.match(rf"^{re.escape(label)}\s*[:\-–]\s*(.+)$", line, re.I)
            if match and not role:
                role = match.group(1).strip()[:200]
                confidence = "high"
        for label in ("company", "organization", "organisation", "employer", "about the company"):
            match = re.match(rf"^{re.escape(label)}\s*[:\-–]\s*(.+)$", line, re.I)
            if match and not company:
                company = match.group(1).strip()[:200]

        looking = re.search(
            r"(?:we are (?:hiring|looking for|seeking)|hiring a[n]?|looking for a[n]?)\s+(.+)$",
            line,
            re.I,
        )
        if looking and not role:
            role = looking.group(1).strip(" .,:;-")[:200]
            confidence = "medium"

    if not role:
        for line in lines[:12]:
            if len(line) > 90 or re.search(r"https?://|www\.|@", line, re.I):
                continue
            if _ROLE_HINT.search(line):
                role = line[:200]
                confidence = "medium"
                break
    if not role and lines:
        first = lines[0]
        if len(first) <= 100 and not re.search(r"https?://|www\.|@", first, re.I):
            role = first[:200]
            confidence = "low"

    if role and company:
        title = f"{role} · {company}"[:200]
    elif role:
        title = role[:200]
    elif company:
        title = f"{company} role"[:200]
    else:
        title = "Job description"

    return {
        "title": title,
        "role_title": role,
        "company": company,
        "confidence": confidence,
    }


_SKILL_CANDIDATES = (
    "python",
    "java",
    "javascript",
    "typescript",
    "sql",
    "react",
    "node.js",
    "nodejs",
    "next.js",
    "nextjs",
    "django",
    "fastapi",
    "spring boot",
    "aws",
    "azure",
    "gcp",
    "docker",
    "kubernetes",
    "git",
    "linux",
    "html",
    "css",
    "mongodb",
    "postgresql",
    "mysql",
    "redis",
    "graphql",
    "rest",
    "power bi",
    "tableau",
    "machine learning",
    "pandas",
    "numpy",
    "c++",
    "c#",
    "go",
    "rust",
    "kotlin",
    "swift",
    "figma",
    "jira",
)


def extract_skill_candidates(text: str, limit: int = 20) -> list[str]:
    """Extract known skill tokens from resume or JD text (deterministic)."""
    haystack = f" {(text or '').lower()} "
    labels = {
        "python": "Python",
        "java": "Java",
        "javascript": "JavaScript",
        "typescript": "TypeScript",
        "sql": "SQL",
        "react": "React",
        "node.js": "Node.js",
        "nodejs": "Node.js",
        "next.js": "Next.js",
        "nextjs": "Next.js",
        "django": "Django",
        "fastapi": "FastAPI",
        "spring boot": "Spring Boot",
        "aws": "AWS",
        "azure": "Azure",
        "gcp": "GCP",
        "docker": "Docker",
        "kubernetes": "Kubernetes",
        "git": "Git",
        "linux": "Linux",
        "html": "HTML",
        "css": "CSS",
        "mongodb": "MongoDB",
        "postgresql": "PostgreSQL",
        "mysql": "MySQL",
        "redis": "Redis",
        "graphql": "GraphQL",
        "rest": "REST",
        "power bi": "Power BI",
        "tableau": "Tableau",
        "machine learning": "Machine Learning",
        "pandas": "Pandas",
        "numpy": "NumPy",
        "c++": "C++",
        "c#": "C#",
        "go": "Go",
        "rust": "Rust",
        "kotlin": "Kotlin",
        "swift": "Swift",
        "figma": "Figma",
        "jira": "Jira",
    }
    found: list[str] = []
    seen: set[str] = set()
    for skill in _SKILL_CANDIDATES:
        token = skill.lower()
        pattern = rf"(?<![a-z0-9+#]){re.escape(token)}(?![a-z0-9+#])"
        if re.search(pattern, haystack):
            label = labels.get(token, skill.title())
            key = label.lower()
            if key not in seen:
                found.append(label)
                seen.add(key)
        if len(found) >= limit:
            break
    return found
