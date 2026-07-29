import io
import uuid
from html import escape
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.auth import CurrentUser
from app.config import Settings
from app.documents import DOCX_MIME, PDF_MIME
from app.errors import ApiError
from app.repository import owned_row


def _sections(structured: dict[str, Any]) -> list[tuple[str, list[str]]]:
    result: list[tuple[str, list[str]]] = []
    for key, value in (structured.get("sections") or {}).items():
        lines = [str(item).strip() for item in value] if isinstance(value, list) else [str(value).strip()]
        result.append((str(key), [line for line in lines if line]))
    return result


def render_docx(structured: dict[str, Any]) -> bytes:
    document = Document()
    unclassified = structured.get("unclassified_blocks") or []
    if unclassified:
        document.add_heading(str(unclassified[0]), level=0)
        for line in unclassified[1:]:
            document.add_paragraph(str(line))
    for section, lines in _sections(structured):
        document.add_heading(section.replace("_", " ").title(), level=1)
        for line in lines:
            document.add_paragraph(line, style="List Bullet" if len(lines) > 1 else None)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(structured: dict[str, Any]) -> bytes:
    output = io.BytesIO()
    styles = getSampleStyleSheet()
    story: list[Any] = []
    unclassified = structured.get("unclassified_blocks") or []
    if unclassified:
        story.append(Paragraph(escape(str(unclassified[0])), styles["Title"]))
        for line in unclassified[1:]:
            story.append(Paragraph(escape(str(line)), styles["BodyText"]))
    for section, lines in _sections(structured):
        story.extend([Spacer(1, 8), Paragraph(escape(section.replace("_", " ").title()), styles["Heading2"])])
        for line in lines:
            story.append(Paragraph(escape(line), styles["BodyText"]))
            story.append(Spacer(1, 4))
    SimpleDocTemplate(output, pagesize=A4, title="Resume", author="").build(story)
    return output.getvalue()


def create_export(
    client, settings: Settings, user: CurrentUser, version_id: str, export_format: str
) -> dict[str, Any]:
    version = owned_row(client, "resume_versions", version_id, user)
    content = (
        render_pdf(version["structured_content"])
        if export_format == "pdf"
        else render_docx(version["structured_content"])
    )
    mime = PDF_MIME if export_format == "pdf" else DOCX_MIME
    export_id = str(uuid.uuid4())
    filename = f"resume-v{version['version_number']}.{export_format}"
    path = f"{user.id}/resumes/{version['resume_id']}/exports/{export_id}/{uuid.uuid4()}.{export_format}"
    try:
        client.storage.from_(settings.document_bucket).upload(
            path, content, {"content-type": mime, "upsert": "false"}
        )
        record = (
            client.table("resume_exports")
            .insert(
                {
                    "id": export_id,
                    "user_id": str(user.id),
                    "resume_version_id": version_id,
                    "export_format": export_format,
                    "storage_path": path,
                    "filename": filename,
                }
            )
            .execute()
            .data[0]
        )
        return record
    except Exception as exc:
        try:
            client.storage.from_(settings.document_bucket).remove([path])
        except Exception:
            pass
        raise ApiError(500, "resume_export_failed", "The resume export could not be created.") from exc


def signed_export(client, settings: Settings, user: CurrentUser, export_id: str) -> dict[str, Any]:
    record = owned_row(client, "resume_exports", export_id, user)
    try:
        response = client.storage.from_(settings.document_bucket).create_signed_url(
            record["storage_path"], settings.export_signed_url_seconds
        )
    except Exception as exc:
        raise ApiError(
            500, "export_download_failed", "A private download link could not be created."
        ) from exc
    url = response.get("signedURL") or response.get("signed_url")
    if not url:
        raise ApiError(500, "export_download_failed", "A private download link could not be created.")
    return {
        "id": record["id"],
        "filename": record["filename"],
        "format": record["export_format"],
        "download_url": url,
        "expires_in": settings.export_signed_url_seconds,
    }
